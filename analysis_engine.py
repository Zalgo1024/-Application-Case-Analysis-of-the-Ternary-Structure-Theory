# 三元结构理论 案例分析引擎 — 分析引擎核心
# Copyright (C) 2026 李政恒 (Li Zhengheng)
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

import json
import re
import os
import sys
import hashlib
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

# Attempt to import pywin32 for Word→PDF conversion
try:
    import win32com.client as win32
    _HAS_PYWIN32 = True
except ImportError:
    _HAS_PYWIN32 = False

# Attempt to import interest network visualizer
try:
    from interest_network_viz import InterestNetworkViz, INTEREST_TYPES, FLOW_TYPES
    _HAS_VIZ = True
except ImportError:
    _HAS_VIZ = False

# Attempt to import image sourcing module
try:
    from image_sourcing import ImageSourcer
    _HAS_IMAGE_SOURCER = True
except ImportError:
    _HAS_IMAGE_SOURCER = False

# Ensure UTF-8 output for Windows terminal
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding='utf-8')


def set_cn_font(run, east_asian_font, western_font='Times New Roman', size=None, bold=False, color=None):
    """设置 run 的中文字体、西文字体、字号、加粗和颜色。可在模块任意位置调用。"""
    run.font.name = western_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asian_font)
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color


class CaseAnalysisEngine:
    def __init__(self, config_path, export_dir=None):
        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = json.load(f)
        self.materials = []
        self.summary_table = []
        if export_dir is None:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            export_dir = os.path.join(base_dir, "报告")
        self.export_dir = export_dir
        if not os.path.exists(self.export_dir):
            os.makedirs(self.export_dir)
        self._figure_counter = 0  # 图表序号计数器
        self._chart_entries = []  # 已嵌入图表记录: [(序号, 标题, 所在章节名)]

    def pre_analyze(self, content):
        """
        Phase 1: Pre-analysis to classify the case and suggest dimensions.
        """
        elements = self.extract_elements(content)
        
        # Determine category based on keywords
        detected_category = "social_governance" # Default
        max_matches = 0
        
        for cat_id, cat_info in self.config.get('categories', {}).items():
            matches = sum(1 for kw in cat_info['keywords'] if kw in content)
            if matches > max_matches:
                max_matches = matches
                detected_category = cat_id
        
        category_name = self.config['categories'][detected_category]['name']
        recommended = self.config['categories'][detected_category]['recommended_dimensions']
        
        return {
            "elements": elements,
            "category_id": detected_category,
            "category_name": category_name,
            "recommended_dimensions": recommended,
            "all_available_dimensions": [
                {"id": d_id, "name": d_info['name']} 
                for d_id, d_info in self.config['dimensions'].items()
            ]
        }

    def run_batch(self, input_list, target_dir=None, analysis_results_list=None, selected_dimensions=None):
        """
        Main entry point for batch processing.
        selected_dimensions: Optional list of dimension IDs to override routing.
        """
        if target_dir:
            self.export_dir = target_dir
            if not os.path.exists(self.export_dir):
                os.makedirs(self.export_dir)

        print("开始执行三元结构理论批量案例分析工作流...\n")
        
        self.materials = self.parse_input(input_list)
        
        for material in self.materials:
            result = self.process_material(material, selected_dimensions)
            self.summary_table.append(result)
            
        self.display_summary()
        self.display_reports(analysis_results_list)
        print(f"\n所有报告已导出至 {os.path.abspath(self.export_dir)} 目录。")

    def parse_input(self, input_list):
        parsed = []
        for item in input_list:
            match = re.match(r"(素材\d+)\s*-\s*(.*)", item, re.DOTALL)
            if match:
                parsed.append({
                    "id": match.group(1),
                    "content": match.group(2),
                    "status": "pending",
                    "type": "N/A"
                })
        return parsed

    def process_material(self, material, selected_dimensions=None):
        elements = self.extract_elements(material['content'])
        material.update(elements)
        
        suitable, reason = self.is_suitable(material)
        if not suitable:
            material['status'] = "拒答"
            material['reason'] = reason
            return material
            
        completeness, missing = self.check_completeness(material)
        if completeness == "failed":
            material['status'] = "信息严重不足"
            material['reason'] = f"缺失项：{'、'.join(missing)}"
            return material
        
        material['report_type'] = "完整版" if completeness == "full" else "精简版"
        material['missing'] = missing
        material['status'] = "成功"
        
        if selected_dimensions:
            material['dimensions'] = selected_dimensions
        else:
            material['dimensions'] = self.route_dimensions(material)
            
        material['is_public_opinion'] = "public_opinion" in material['dimensions']
        
        return material

    def extract_elements(self, content):
        return {
            "subject": "待提取", 
            "event": "待提取", 
            "background": "待提取"
        }

    def is_suitable(self, material):
        content = material['content'].lower()
        
        science_keywords = ["物理", "化学", "数学", "医疗", "地质", "天体"]
        if any(kw in content for kw in science_keywords):
            return False, "该内容属于自然科学范畴，三元结构理论仅适用于人类社会行为与现象分析"
        
        is_po = any(kw in content for kw in self.config['dimensions']['public_opinion']['routing_keywords'])
        if is_po:
            if "ai生成的" in content or "虚构的舆情" in content:
                return False, "虚构的舆情情节或 AI 生成的虚拟舆情除非明确要求分析虚构世界逻辑，否则不予分析"

        social_keywords = ["公司", "组织", "群体", "事件", "制度", "文化", "博弈", "行为", "舆情", "舆论", "传播", "热议"]
        if not any(kw in content for kw in social_keywords):
            return False, "该内容不涉及明确的人类社会行为或社会现象"

        return True, ""

    def check_completeness(self, material):
        missing = []
        is_po = any(kw in material['content'].lower() for kw in self.config['dimensions']['public_opinion']['routing_keywords'])
        
        if not material.get('subject') or material.get('subject') == "待提取": 
            missing.append("主体 (舆情核心主体/传播主体)" if is_po else "主体")
        if not material.get('event') or material.get('event') == "待提取": 
            missing.append("事件/现象 (舆情议题/关键行为)" if is_po else "事件/现象")
        if not material.get('background') or material.get('background') == "待提取": 
            missing.append("时间/背景 (发酵节点/社会情绪)" if is_po else "时间/背景")
        
        if len(missing) >= 2:
            return "failed", missing
        if len(missing) == 1:
            return "partial", missing
        return "full", []

    def route_dimensions(self, material):
        content = material['content'].lower()
        
        po_keywords = self.config['dimensions']['public_opinion']['routing_keywords']
        if any(kw in content for kw in po_keywords):
            return ["public_opinion"]
            
        selected = []
        for dim_id, dim_info in self.config['dimensions'].items():
            if dim_id == "public_opinion": continue
            if any(kw in content for kw in dim_info['routing_keywords']):
                selected.append(dim_id)
        
        return selected[:2] if selected else ["three_point"]

    def display_summary(self):
        print("| 素材 ID | 处理状态 | 核心信息 | 异常原因 / 信息缺口 | 报告类型 |")
        print("| --- | --- | --- | --- | --- |")
        for m in self.summary_table:
            info = f"{m.get('subject', 'N/A')} + {m.get('event', 'N/A')}"
            print(f"| {m['id']} | {m['status']} | {info} | {m.get('reason', '无')} | {m.get('report_type', '-')} |")
        print("\n" + "="*50 + "\n")

    def generate_report_data(self, m, analysis_results=None):
        """
        Generates structured data for the report (Updated for v2.2 high tension narrative).
        """
        if analysis_results is None:
            analysis_results = {}

        data = {
            "title": analysis_results.get("title", f"【{m['id']}】{m.get('subject', '案例')}分析 —— 基于三元结构理论"),
            "core_tension": analysis_results.get("core_tension", "案例核心冲突点：待补充。"),
            "copyright": self.config['copyright'],
            "is_brief": m['report_type'] == "精简版",
            "missing": m.get('missing', []),
            "abstract": analysis_results.get("abstract", "[此处生成摘要，侧重：背景→问题→维度→结论概要...]"),
            "sections": [
                {
                    "name": "1. 案例背景与概述",
                    "content": analysis_results.get("background", f"时间/背景：{m.get('background')}\n主体：{m.get('subject')}\n事件/现象：{m.get('event')}")
                },
                {
                    "name": "2. 核心问题定义",
                    "content": analysis_results.get("core_problem", "[此处定义 1-2 个具体核心矛盾，强化影响分析...]")
                },
                {
                    "name": "3. 基于三元结构理论的分析",
                    "sub_sections": []
                },
                {
                    "name": "4. 解决方案与实施建议",
                    "content": analysis_results.get("solutions", "[提供方案，明确能缓解/规避的具体影响...]")
                },
                {
                    "name": "5. 结论",
                    "content": analysis_results.get("conclusion", "[本质机制提炼 + 趋势判断...]"),
                    "open_question": analysis_results.get("open_question", "留给未来的问题是：______")
                }
            ],
            "reference": self.config['reference']
        }
        
        if m.get('is_public_opinion'):
            # Specific PO sub-sections
            data['sections'][2]['sub_sections'] = [
                {"name": "3.1 生存维度：舆情的存在根基与核心受众", "content": analysis_results.get("po_survival", "[分析舆情存在根基...]")},
                {"name": "3.2 繁衍维度：舆情的传播特征与演化规律", "content": analysis_results.get("po_reproduction", "[分析复制/变异性...]")},
                {"name": "3.3 逆反维度：舆情引发的逆反行为与连锁影响", "content": analysis_results.get("po_reactivity", "[分析逆反性质与触发机制...]")},
                {"name": "3.4 三维度交叉分析", "content": analysis_results.get("po_cross_analysis", "[分析联动放大效应...]")},
                {"name": "视角切换/推演", "content": analysis_results.get("perspective_switch", "[视角切换段落...]"), "is_special": True}
            ]
        else:
            for dim in m['dimensions']:
                dim_name = self.config['dimensions'][dim]['name']
                content_key = f"dim_{dim}"
                data['sections'][2]['sub_sections'].append({
                    "name": f"3.{len(data['sections'][2]['sub_sections'])+1} {dim_name}维度分析",
                    "content": analysis_results.get(content_key, f"[使用 {dim_name} 维度进行拆解...]")
                })
            # Add perspective switch as a separate section or part of sections
            if analysis_results.get("perspective_switch"):
                data['sections'][2]['sub_sections'].append({
                    "name": "视角切换/反事实推演",
                    "content": analysis_results.get("perspective_switch"),
                    "is_special": True
                })
            
        return data

    def display_reports(self, analysis_results_list=None):
        for i, m in enumerate(self.summary_table):
            if m['status'] == "成功":
                analysis_results = analysis_results_list[i] if analysis_results_list and i < len(analysis_results_list) else None
                report_content = self.generate_report_data(m, analysis_results)
                self.print_report(report_content)
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_subject = re.sub(r'[\\/*?:"<>|]', '_', m.get('subject', '案例'))
                report_folder_name = f"{safe_subject}_{timestamp}"
                report_path = os.path.join(self.export_dir, report_folder_name)
                
                if not os.path.exists(report_path):
                    os.makedirs(report_path)
                
                self.export_to_word(m, report_content, report_path)
                self.export_to_pdf(m, report_content, report_path)

    def print_report(self, data):
        print(f"\n## {data['title']}")
        print(f"\n> {data['copyright']}\n")
        if data['is_brief']:
            print(f"**注意：以下分析基于有限信息，缺失项：{'、'.join(data['missing'])}。**\n")
        for sec in data['sections']:
            print(f"\n### {sec['name']}")
            if 'content' in sec:
                print(sec['content'])
            if 'sub_sections' in sec:
                for sub in sec['sub_sections']:
                    print(f"#### {sub['name']}\n{sub['content']}")
        print(f"\n**参考文献**：{data['reference']}")
        print("-" * 30)

    def export_to_word(self, m, data, report_path):
        doc = Document()
        
        # Set Page Margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Title
        title_para = doc.add_heading(data['title'], 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_para.runs:
            run.font.color.rgb = RGBColor(44, 62, 80)

        # Core Tension (Quote Block style)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(f"「{data['core_tension']}」")
        run.italic = True
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(52, 152, 219) # Blue

        # Copyright & Timestamp
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{data['copyright']}\n生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(127, 140, 141)

        if data['is_brief']:
            p = doc.add_paragraph()
            run = p.add_run(f"【精简版说明】缺失项：{'、'.join(data['missing'])}。")
            run.bold = True
            run.font.color.rgb = RGBColor(192, 57, 43)

        # Abstract
        doc.add_heading('摘要', level=1)
        p = doc.add_paragraph()
        run = p.add_run(data['abstract'])
        run.font.size = Pt(10.5)

        # Sections
        for sec in data['sections']:
            h = doc.add_heading(sec['name'], level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(52, 73, 94)

            if 'content' in sec and sec['content']:
                # Support rhythmic short paragraphs (split by newline)
                paras = sec['content'].split('\n')
                for para_text in paras:
                    if not para_text.strip(): continue
                    p = doc.add_paragraph()
                    run = p.add_run(para_text.strip())
                    # If very short, treat as rhythmic pause
                    if len(para_text.strip()) < 20:
                        run.bold = True
                        p.paragraph_format.space_before = Pt(6)
                    run.font.size = Pt(11)
            
            if 'sub_sections' in sec:
                for sub in sec['sub_sections']:
                    h2 = doc.add_heading(sub['name'], level=2)
                    for run in h2.runs:
                        run.font.color.rgb = RGBColor(41, 128, 185)
                    if sub.get('content'):
                        paras = sub['content'].split('\n')
                        for para_text in paras:
                            if not para_text.strip(): continue
                            p = doc.add_paragraph()
                            run = p.add_run(para_text.strip())
                            if len(para_text.strip()) < 20:
                                run.bold = True
                                p.paragraph_format.space_before = Pt(6)
                            run.font.size = Pt(11)
            
            # Special logic for Open Question in Conclusion
            if sec['name'] == "5. 结论" and sec.get('open_question'):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                run = p.add_run(sec['open_question'])
                run.bold = True
                run.italic = True
                run.font.color.rgb = RGBColor(44, 62, 80)

        # Reference
        doc.add_paragraph("\n")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"--- 参考文献 ---\n{data['reference']}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(149, 165, 166)
        
        file_path = os.path.join(report_path, f"{m['id']}_分析报告.docx")
        doc.save(file_path)

    def export_to_pdf(self, m, data, report_path):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        font_paths = ["C:\\Windows\\Fonts\\simhei.ttf", "C:\\Windows\\Fonts\\simsun.ttc", "C:\\Windows\\Fonts\\msyh.ttc", "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf"]
        font_loaded = False
        font_name = "Chinese"
        for path in font_paths:
            if os.path.exists(path):
                pdf.add_font(font_name, "", path)
                pdf.add_font(font_name, "B", path)
                pdf.set_font(font_name, size=12)
                font_loaded = True
                break
        
        if not font_loaded:
            pdf.set_font("Arial", size=12)
            font_name = "Arial"

        def encode_str(s): return s

        # Top line
        pdf.set_draw_color(41, 128, 185)
        pdf.set_line_width(1)
        pdf.line(10, 10, 200, 10)

        # Title
        pdf.ln(10)
        pdf.set_font(font_name, size=20)
        pdf.set_text_color(44, 62, 80)
        pdf.multi_cell(0, 15, encode_str(data['title']), align='C')
        
        # Core Tension Block
        pdf.ln(5)
        pdf.set_fill_color(244, 246, 247)
        pdf.set_text_color(52, 152, 219)
        pdf.set_font(font_name, "B", size=12)
        pdf.multi_cell(0, 10, encode_str(f"「{data['core_tension']}」"), align='C', fill=True)
        
        # Copyright & Time
        pdf.ln(2)
        pdf.set_font(font_name, size=8)
        pdf.set_text_color(127, 140, 141)
        pdf.cell(0, 5, encode_str(data['copyright']), ln=True, align='R')
        pdf.cell(0, 5, encode_str(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"), ln=True, align='R')
        
        if data['is_brief']:
            pdf.ln(5)
            pdf.set_font(font_name, size=10)
            pdf.set_text_color(192, 57, 43)
            pdf.multi_cell(0, 8, encode_str(f"【精简版说明】缺失项：{'、'.join(data['missing'])}。"))

        # Abstract
        pdf.ln(5)
        pdf.set_fill_color(236, 240, 241)
        pdf.set_font(font_name, size=14)
        pdf.set_text_color(52, 73, 94)
        pdf.cell(0, 12, encode_str("摘要"), ln=True, fill=True)
        pdf.set_font(font_name, size=11)
        pdf.set_text_color(0, 0, 0)
        pdf.multi_cell(0, 8, encode_str(data['abstract']))

        # Sections
        for sec in data['sections']:
            pdf.ln(8)
            pdf.set_fill_color(236, 240, 241)
            pdf.set_font(font_name, size=14)
            pdf.set_text_color(44, 62, 80)
            pdf.cell(0, 12, encode_str(sec['name']), ln=True, fill=True)
            pdf.set_font(font_name, size=11)
            pdf.set_text_color(0, 0, 0)
            
            if 'content' in sec and sec['content']:
                paras = sec['content'].split('\n')
                for p_text in paras:
                    p_text = p_text.strip()
                    if not p_text: continue
                    pdf.set_x(10) # Ensure we start from left margin
                    if len(p_text) < 20: # Rhythmic pause
                        pdf.set_font(font_name, "B", size=11)
                        pdf.multi_cell(0, 8, encode_str(p_text))
                        pdf.set_font(font_name, size=11)
                    else:
                        pdf.multi_cell(0, 8, encode_str(p_text))
            
            if 'sub_sections' in sec:
                for sub in sec['sub_sections']:
                    pdf.ln(4)
                    pdf.set_x(10)
                    pdf.set_font(font_name, size=12)
                    pdf.set_text_color(41, 128, 185)
                    pdf.cell(0, 10, encode_str(sub['name']), ln=True)
                    pdf.set_font(font_name, size=11)
                    pdf.set_text_color(0, 0, 0)
                    if sub.get('content'):
                        paras = sub['content'].split('\n')
                        for p_text in paras:
                            p_text = p_text.strip()
                            if not p_text: continue
                            pdf.set_x(10)
                            if len(p_text) < 20: # Rhythmic pause
                                pdf.set_font(font_name, "B", size=11)
                                pdf.multi_cell(0, 8, encode_str(p_text))
                                pdf.set_font(font_name, size=11)
                            else:
                                pdf.multi_cell(0, 8, encode_str(p_text))
            
            if sec['name'] == "5. 结论" and sec.get('open_question'):
                pdf.ln(5)
                pdf.set_x(10)
                pdf.set_font(font_name, "B", size=11)
                pdf.set_text_color(44, 62, 80)
                pdf.multi_cell(0, 8, encode_str(sec['open_question']))

        # Reference
        pdf.ln(15)
        pdf.set_draw_color(189, 195, 199)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)
        pdf.set_font(font_name, size=8)
        pdf.set_text_color(149, 165, 166)
        pdf.multi_cell(0, 6, encode_str(f"参考文献：{data['reference']}"), align='C')

        file_path = os.path.join(report_path, f"{m['id']}_分析报告.pdf")
        try:
            pdf.output(file_path)
        except Exception as e:
            print(f"PDF 导出失败: {e}")

    @staticmethod
    def _clean_ai_text(text):
        """清理 AI 写作格式痕迹"""
        if not isinstance(text, str) or not text:
            return text
        text = re.sub(r'—{3,}', '——', text)
        return text

    # ====================================================================
    #  v4.0 文本导出管线：六部分固定结构 + 示意图 + 引用块 + 结论句
    # ====================================================================

    def export_from_text(self, title, markdown_text, export_dir=None, force_new=False):
        """
        直接从完整 Markdown 分析文本导出 Word + PDF 报告。
        支持 v4.0 六部分固定结构：封面页→框架前置说明→事实摘要→分析正文→结论→附录。

        v6.1.1：默认启用内容哈希去重——相同 title + markdown_text 不会重复生成。
        设置 force_new=True 可强制创建新报告。

        Args:
            title: 报告标题
            markdown_text: 完整的 Markdown 格式分析报告文本
            export_dir: 可选，覆盖默认导出目录
            force_new: 强制创建新报告（默认 False，启用哈希去重）

        Returns:
            dict: {"word": "文件路径", "pdf": "文件路径", "folder": "文件夹路径"}
        """
        target_dir = export_dir or self.export_dir
        if not os.path.exists(target_dir):
            os.makedirs(target_dir)

        # ---- v6.1.1：内容哈希去重 ----
        content_hash = hashlib.sha256(f"{title}{markdown_text}".encode('utf-8')).hexdigest()[:16]
        safe_title = re.sub(r'[\\/*?:"<>|]', '_', title)[:50]

        if not force_new:
            # 扫描已有报告文件夹，查找匹配的 .report_hash
            try:
                for entry in os.scandir(target_dir):
                    if entry.is_dir():
                        hash_path = os.path.join(entry.path, '.report_hash')
                        if os.path.isfile(hash_path):
                            try:
                                with open(hash_path, 'r', encoding='utf-8') as hf:
                                    stored_hash = hf.read().strip()
                                if stored_hash == content_hash:
                                    # 找到已有报告 —— 检查关键文件是否存在
                                    word_path = os.path.join(entry.path, f"{safe_title}.docx")
                                    pdf_path = os.path.join(entry.path, f"{safe_title}.pdf")
                                    if os.path.isfile(word_path) and os.path.isfile(pdf_path):
                                        print(f"  [去重] 内容未变，使用已有报告：{entry.path}")
                                        return {
                                            "folder": entry.path,
                                            "word": word_path,
                                            "pdf": pdf_path,
                                        }
                                    else:
                                        # 哈希匹配但文件缺失 → 重新生成到此文件夹
                                        report_path = entry.path
                                        break
                            except Exception:
                                pass  # 哈希文件损坏，忽略
                else:
                    report_path = None  # 未找到匹配
            except Exception:
                report_path = None  # 扫描失败，降级到新建
        else:
            report_path = None  # force_new 模式

        # ---- 创建新文件夹（如有需要） ----
        if report_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_folder_name = f"{safe_title}_{timestamp}"
            report_path = os.path.join(target_dir, report_folder_name)
            if not os.path.exists(report_path):
                os.makedirs(report_path)

            # 写入哈希文件
            hash_path = os.path.join(report_path, '.report_hash')
            with open(hash_path, 'w', encoding='utf-8') as hf:
                hf.write(content_hash)

        data = self._parse_markdown_report(title, markdown_text)
        pseudo_material = {"id": safe_title}

        word_path = os.path.join(report_path, f"{safe_title}.docx")
        self._export_word_from_text(pseudo_material, data, report_path)
        pdf_path = self._convert_word_to_pdf(word_path)

        result = {
            "folder": report_path,
            "word": word_path,
            "pdf": pdf_path
        }
        print(f"\n报告已导出：\n  Word: {result['word']}\n  PDF:  {result['pdf']}")
        return result

    def _parse_markdown_report(self, title, markdown_text):
        """
        将 Markdown 格式的分析报告解析为结构化数据。
        支持六部分固定结构：封面页→框架前置说明→事实摘要→分析正文→结论→附录。
        保留代码块（纯文本示意图）、引用块（>）、结论句（→/⇒）标记。
        """
        lines = markdown_text.split('\n')
        data = {
            "title": title,
            "copyright": self.config['copyright'],
            "framework_preamble": "",
            "fact_summary": "",
            "sections": [],
            "reference": self.config['reference'],
            "is_brief": False,
            "missing": []
        }

        current_section = None
        current_sub_section = None
        current_content_lines = []
        in_code_block = False
        code_block_lines = []

        PREAMBLE_NAMES = {"分析框架前置说明", "框架前置说明", "框架说明"}
        FACT_SUMMARY_NAMES = {"案例事实摘要", "事实摘要", "案例摘要"}

        def flush_content():
            text = '\n'.join(current_content_lines).strip()
            if not text:
                return
            if current_sub_section is not None:
                current_sub_section['content'] = text
            elif current_section is not None:
                current_section['content'] = text

        for line in lines:
            stripped = line.strip()

            # 代码块处理：``` 开始/结束
            if stripped.startswith('```'):
                if in_code_block:
                    in_code_block = False
                    code_content = '\n'.join(code_block_lines)
                    current_content_lines.append(f'\n```DIAGRAM\n{code_content}\n```\n')
                    code_block_lines = []
                else:
                    in_code_block = True
                    code_block_lines = []
                continue

            if in_code_block:
                code_block_lines.append(line.rstrip())
                continue

            if not stripped or stripped == '---':
                continue

            if stripped.startswith('#### '):
                flush_content()
                current_content_lines = []
                heading_text = stripped[5:].strip()
                current_content_lines.append(heading_text)
                continue

            if stripped.startswith('### '):
                flush_content()
                current_content_lines = []
                sub_name = stripped[4:].strip()
                current_sub_section = {"name": sub_name, "content": ""}
                if current_section is not None:
                    if 'sub_sections' not in current_section:
                        current_section['sub_sections'] = []
                    current_section['sub_sections'].append(current_sub_section)
                continue

            if stripped.startswith('## '):
                flush_content()
                current_content_lines = []
                sec_name = stripped[3:].strip()

                if sec_name in PREAMBLE_NAMES:
                    current_section = {"name": sec_name, "content": "", "is_preamble": True}
                    data['sections'].append(current_section)
                    current_sub_section = None
                    continue

                if sec_name in FACT_SUMMARY_NAMES:
                    current_section = {"name": sec_name, "content": "", "is_fact_summary": True}
                    data['sections'].append(current_section)
                    current_sub_section = None
                    continue

                current_section = {"name": sec_name, "content": ""}
                data['sections'].append(current_section)
                current_sub_section = None
                continue

            # 标题前的文本跳过（封面页信息由引擎自动生成）
            if current_section is None and current_sub_section is None:
                continue

            # 保留 > 引用标记和 →/⇒ 结论标记（不再 strip > 前缀）
            line_text = stripped
            if line_text:
                if current_section and current_section.get('is_preamble'):
                    data['framework_preamble'] += line_text + '\n'
                if current_section and current_section.get('is_fact_summary'):
                    data['fact_summary'] += line_text + '\n'
                current_content_lines.append(line_text)

        flush_content()

        data['framework_preamble'] = data['framework_preamble'].strip()
        data['fact_summary'] = data['fact_summary'].strip()

        return data

    # ---------- Word 导出（v4.0 六部分结构） ----------

    def _export_word_from_text(self, m, data, report_path):
        """从文本解析数据导出 Word 文档（六部分固定结构：封面→框架→事实→分析→结论→附录）"""
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(3.7)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)

        # ========== 配色方案 ==========
        COLOR_TITLE    = RGBColor(0x1B, 0x3A, 0x5C)
        COLOR_H1       = RGBColor(0x1B, 0x3A, 0x5C)
        COLOR_H2       = RGBColor(0x2C, 0x5F, 0x8A)
        COLOR_BODY     = RGBColor(0x33, 0x33, 0x33)
        COLOR_ACCENT   = RGBColor(0xC0, 0x39, 0x2B)
        COLOR_COPYRIGHT = RGBColor(0x7F, 0x8C, 0x8D)
        COLOR_REF      = RGBColor(0x5D, 0x6D, 0x7E)
        COLOR_QUOTE    = RGBColor(0x5D, 0x6D, 0x7E)
        COLOR_CONCLUSION = RGBColor(0xC0, 0x39, 0x2B)
        COLOR_DIAGRAM_LABEL = RGBColor(0x1B, 0x3A, 0x5C)

        # ========== 字号方案 ==========
        SZ_TITLE   = Pt(24)
        SZ_H1      = Pt(16)
        SZ_H2      = Pt(14)
        SZ_BODY    = Pt(12)
        SZ_SMALL   = Pt(10.5)
        SZ_DIAGRAM = Pt(10.5)
        LS_TITLE   = Pt(36)
        LS_H1      = Pt(28)
        LS_H2      = Pt(24)
        LS_BODY    = Pt(22)
        INDENT_BODY = Pt(24)

        def add_rich_para_line(doc, line, indent=True, font_size=SZ_BODY, line_spacing=LS_BODY,
                               font_name='仿宋', color=COLOR_BODY):
            """添加单行正文段落，支持 **加粗** 标记"""
            line = self._clean_ai_text(line.strip())
            if not line:
                return
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = line_spacing
            if indent:
                p.paragraph_format.first_line_indent = INDENT_BODY
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_cn_font(run, font_name, size=font_size, bold=True, color=color)
                elif part:
                    run = p.add_run(part)
                    set_cn_font(run, font_name, size=font_size, color=color)

        def add_quote_para(doc, text):
            """添加引用块段落（缩进+灰字+左侧竖线标记）"""
            text = text.strip()
            if not text:
                return
            if text.startswith('> '):
                text = text[2:]
            elif text.startswith('>'):
                text = text[1:]
            text = text.strip()
            if not text:
                return
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = LS_BODY
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Pt(0)
            run_bar = p.add_run('▎ ')
            set_cn_font(run_bar, '仿宋', size=SZ_BODY, color=COLOR_QUOTE)
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_cn_font(run, '楷体', size=SZ_BODY, bold=True, color=COLOR_QUOTE)
                elif part:
                    run = p.add_run(part)
                    set_cn_font(run, '楷体', size=SZ_BODY, color=COLOR_QUOTE)

        def add_conclusion_para(doc, text):
            """添加结论句段落（→/⇒ 开头，暖红色加粗）"""
            text = text.strip()
            if not text:
                return
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = LS_BODY
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(6)
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_cn_font(run, '仿宋', size=SZ_BODY, bold=True, color=COLOR_CONCLUSION)
                elif part:
                    run = p.add_run(part)
                    set_cn_font(run, '仿宋', size=SZ_BODY, bold=True, color=COLOR_CONCLUSION)

        def add_list_item(doc, text, indent=True):
            """添加列表项段落（圆点标记 + 正文）"""
            text = text.strip()
            if not text:
                return
            if text.startswith('- '):
                text = text[2:].strip()
            elif text.startswith('-'):
                text = text[1:].strip()
            if not text:
                return
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = LS_BODY
            if indent:
                p.paragraph_format.first_line_indent = INDENT_BODY
            run = p.add_run('• ')
            set_cn_font(run, '仿宋', size=SZ_BODY, color=COLOR_BODY)
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_cn_font(run, '仿宋', size=SZ_BODY, bold=True, color=COLOR_BODY)
                elif part:
                    run = p.add_run(part)
                    set_cn_font(run, '仿宋', size=SZ_BODY, color=COLOR_BODY)

        def add_diagram_block(doc, diagram_text):
            """
            添加示意图块。
            优先尝试解析为可视化网络数据（JSON格式），生成SVG/PNG图表嵌入。
            回退到纯文本字符画模式。
            """
            diagram_text = diagram_text.strip()
            if not diagram_text:
                return

            # 尝试解析为可视化网络 JSON
            viz_image = self._try_render_network_viz(doc, diagram_text, report_dir=report_path)
            if viz_image:
                return

            # 回退：纯文本字符画
            diag_lines = diagram_text.split('\n')
            for line in diag_lines:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.first_line_indent = Pt(0)
                # 标签行【XX：XX】用蓝色加粗
                if line.strip().startswith('【') and '】' in line:
                    run = p.add_run(line)
                    set_cn_font(run, '黑体', western_font='Consolas', size=SZ_DIAGRAM, bold=True, color=COLOR_DIAGRAM_LABEL)
                else:
                    run = p.add_run(line)
                    set_cn_font(run, '仿宋', western_font='Consolas', size=SZ_DIAGRAM, color=COLOR_BODY)

        def embed_image_with_attribution(doc, img_path, attribution_text):
            """
            在 Word 中嵌入外部图片（网络扒图）并添加来源标注。

            Args:
                doc: python-docx Document
                img_path: 图片本地绝对路径
                attribution_text: 来源标注文本（如 "Wikimedia Commons | CC BY-SA 4.0 | 原始链接"）
            """
            if not img_path or not os.path.exists(img_path):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(0)
                run = p.add_run(f"[图片未找到: {img_path}]")
                set_cn_font(run, '仿宋', size=SZ_SMALL, color=COLOR_REF)
                return

            # 嵌入图片（居中）
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            p_img.paragraph_format.first_line_indent = Pt(0)

            try:
                run_img = p_img.add_run()
                run_img.add_picture(img_path, width=Inches(5.5))
            except Exception as e:
                run_fail = p_img.add_run(f"[无法嵌入图片: {e}]")
                set_cn_font(run_fail, '仿宋', size=SZ_SMALL, color=COLOR_ACCENT)

            # 来源标注（居中、小号灰色）
            if attribution_text:
                p_attr = doc.add_paragraph()
                p_attr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_attr.paragraph_format.space_after = Pt(10)
                p_attr.paragraph_format.first_line_indent = Pt(0)
                run_attr = p_attr.add_run(f"[来源：{attribution_text}]")
                set_cn_font(run_attr, '仿宋', size=Pt(9), color=COLOR_REF)

        def render_content(doc, content_text, indent=True):
            """渲染内容文本，自动识别引用块、示意图、结论句、列表、普通段落
            关键修复：先用正则提取 DIAGRAM 块再分段，防止空行拆散示意图"""
            content_text = content_text.strip()
            if not content_text:
                return

            # Step 1: 提取 DIAGRAM 块，防止被 \n\n 拆散
            diagram_map = {}
            diag_counter = [0]

            def extract_diagram(match):
                key = f'__DIAGRAM_{diag_counter[0]}__'
                diagram_map[key] = match.group(1).strip()
                diag_counter[0] += 1
                return f'\n\n{key}\n\n'

            processed = re.sub(
                r'```DIAGRAM\n(.*?)```',
                extract_diagram,
                content_text,
                flags=re.DOTALL
            )

            # Step 2: 按双换行分段
            paragraphs = re.split(r'\n\s*\n', processed)

            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # 示意图占位符
                if para.startswith('__DIAGRAM_') and para.endswith('__') and para in diagram_map:
                    add_diagram_block(doc, diagram_map[para])
                    continue

                # 外部图片嵌入（整段为 ![attribution](path) 格式）
                img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', para)
                if img_match:
                    embed_image_with_attribution(doc, img_match.group(2), img_match.group(1))
                    continue

                # 引用块（整段以 > 开头）
                if para.startswith('> ') or (para.startswith('>') and not para.startswith('>>')):
                    add_quote_para(doc, para)
                    continue

                # 结论句（整段以 →/⇒ 开头）
                if para.startswith('→ ') or para.startswith('⇒ ') or para == '→' or para == '⇒':
                    add_conclusion_para(doc, para)
                    continue

                # 逐行判断混合内容
                for line in para.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith('> ') or (line.startswith('>') and not line.startswith('>>')):
                        add_quote_para(doc, line)
                    elif line.startswith('→ ') or line.startswith('⇒ ') or line == '→' or line == '⇒':
                        add_conclusion_para(doc, line)
                    elif line.startswith('- '):
                        add_list_item(doc, line, indent=indent)
                    elif re.match(r'^!\[[^\]]*\]\([^)]+\)$', line):
                        # 行内图片嵌入 ![alt](path)
                        m = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', line)
                        embed_image_with_attribution(doc, m.group(2), m.group(1))
                    else:
                        add_rich_para_line(doc, line, indent=indent)

        # ========== 一、封面页 ==========
        title_para = doc.add_heading(data['title'], 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        title_para.paragraph_format.line_spacing = LS_TITLE
        title_para.paragraph_format.space_after = Pt(12)
        for run in title_para.runs:
            set_cn_font(run, '华文中宋', size=SZ_TITLE, color=COLOR_TITLE)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(20)
        run = p.add_run(data['copyright'])
        set_cn_font(run, '仿宋', size=SZ_SMALL, color=COLOR_COPYRIGHT)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p2.paragraph_format.line_spacing = Pt(20)
        run2 = p2.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        set_cn_font(run2, '仿宋', size=SZ_SMALL, color=COLOR_COPYRIGHT)

        if data.get('is_brief') and data.get('missing'):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = LS_BODY
            run = p.add_run(f"【精简版说明】缺失项：{'、'.join(data['missing'])}。")
            set_cn_font(run, '黑体', size=SZ_BODY, color=COLOR_ACCENT)

        # 封面页分节：新节（封面无页眉页脚，正文有页眉页脚）
        new_section = doc.add_section()
        new_section.top_margin = Cm(3.7)
        new_section.bottom_margin = Cm(3.5)
        new_section.left_margin = Cm(2.8)
        new_section.right_margin = Cm(2.6)

        # ---- 页眉：居中案例名称（楷体小号灰字） ----
        header = new_section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hrun = hp.add_run(data['title'])
        set_cn_font(hrun, '楷体', size=Pt(9), color=COLOR_REF)
        # 页眉底部细线
        hp.paragraph_format.space_after = Pt(2)

        # ---- 页脚：居中页码（1/17 格式） ----
        footer = new_section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # PAGE 域
        fld_begin1 = OxmlElement('w:fldChar')
        fld_begin1.set(qn('w:fldCharType'), 'begin')
        fp.add_run()._r.append(fld_begin1)

        instr1 = OxmlElement('w:instrText')
        instr1.set(qn('xml:space'), 'preserve')
        instr1.text = ' PAGE '
        fp.add_run()._r.append(instr1)

        fld_end1 = OxmlElement('w:fldChar')
        fld_end1.set(qn('w:fldCharType'), 'end')
        fp.add_run()._r.append(fld_end1)

        # 分隔符 /
        run_slash = fp.add_run(' / ')
        set_cn_font(run_slash, '仿宋', size=Pt(9), color=COLOR_REF)

        # NUMPAGES 域
        fld_begin2 = OxmlElement('w:fldChar')
        fld_begin2.set(qn('w:fldCharType'), 'begin')
        fp.add_run()._r.append(fld_begin2)

        instr2 = OxmlElement('w:instrText')
        instr2.set(qn('xml:space'), 'preserve')
        instr2.text = ' NUMPAGES '
        fp.add_run()._r.append(instr2)

        fld_end2 = OxmlElement('w:fldChar')
        fld_end2.set(qn('w:fldCharType'), 'end')
        fp.add_run()._r.append(fld_end2)

        # ========== 二~六、正文部分 ==========
        for sec in data['sections']:
            h = doc.add_heading(sec['name'], level=1)
            h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            h.paragraph_format.line_spacing = LS_H1
            h.paragraph_format.space_before = Pt(12)
            for run in h.runs:
                set_cn_font(run, '黑体', size=SZ_H1, color=COLOR_H1)

            if sec.get('content'):
                render_content(doc, sec['content'],
                               indent=not sec.get('is_preamble'))

            if sec.get('sub_sections'):
                for sub in sec['sub_sections']:
                    h2 = doc.add_heading(sub['name'], level=2)
                    h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    h2.paragraph_format.line_spacing = LS_H2
                    h2.paragraph_format.space_before = Pt(8)
                    for run in h2.runs:
                        set_cn_font(run, '楷体', size=SZ_H2, color=COLOR_H2)
                    if sub.get('content'):
                        render_content(doc, sub['content'])

        # ========== 参考文献（支持超链接） ==========
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run("参考文献")
        set_cn_font(run, '仿宋', size=SZ_BODY, bold=True, color=COLOR_REF)

        # 优先使用附录 section 中的实际数据来源，而非 config 默认值
        appendix_ref_text = ''
        for sec in data['sections']:
            if sec['name'] == '附录':
                content = sec.get('content', '')
                # 提取 **数据来源**： 和 **分析框架** 之间的内容
                in_ref = False
                for line in content.split('\n'):
                    if '**数据来源**' in line or '**数据来源**：' in line:
                        in_ref = True
                        continue
                    if '**分析框架**' in line:
                        break
                    if in_ref:
                        appendix_ref_text += line + '\n'
                break

        ref_text = appendix_ref_text.strip() or data['reference']
        ref_lines = ref_text.strip().split('\n')
        for rline in ref_lines:
            rline = rline.strip()
            if not rline:
                continue
            rp = doc.add_paragraph()
            rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            rp.paragraph_format.line_spacing = Pt(18)
            self._render_reference_line(rp, rline, SZ_SMALL, COLOR_REF)

        # ========== 自动图表目录（如报告已手动写入则不重复） ==========
        if self._chart_entries:
            # 检查文档中是否已存在 "图表目录" 标题（避免重复插入）
            has_manual = False
            for para in doc.paragraphs:
                if '图表目录' in (para.text or ''):
                    has_manual = True
                    break
            if not has_manual:
                doc.add_paragraph()  # 间距
                h_dir = doc.add_heading('图表目录', level=2)
                for run in h_dir.runs:
                    set_cn_font(run, '楷体', size=SZ_H2, color=COLOR_H2)
                for fig_num, fig_title in self._chart_entries:
                    ep = doc.add_paragraph()
                    ep.paragraph_format.first_line_indent = Cm(0.74)
                    er = ep.add_run(f'图{fig_num}：{fig_title}')
                    set_cn_font(er, '仿宋', size=SZ_BODY, color=COLOR_BODY)

        file_path = os.path.join(report_path, f"{m['id']}.docx")
        doc.save(file_path)

    def _try_render_network_viz(self, doc, diagram_text, report_dir=None):
        """
        尝试将 diagram_text 解析为利益关系网络 JSON 并生成可视化图表嵌入 Word。

        支持的 JSON 格式：
        ```DIAGRAM
        {"viz": "network", "title": "...", "nodes": [...], "edges": [...]}
        ```

        节点格式：{"id": "A", "label": "...", "type": "political", "detail": "..."}
        边格式：{"source": "A", "target": "B", "label": "...", "type": "economic"}

        Args:
            doc: python-docx Document 对象
            diagram_text: DIAGRAM 代码块文本
            report_dir: 报告输出目录（用于保存交互式 HTML），默认回退到 export_dir

        返回 True 如果成功渲染，否则返回 False。
        """
        if not _HAS_VIZ:
            return False

        # 尝试提取 JSON
        json_match = re.search(r'\{.*"viz"\s*:\s*"network".*\}', diagram_text, re.DOTALL)
        if not json_match:
            return False

        try:
            data = json.loads(json_match.group())
            if data.get('viz') != 'network':
                return False

            viz = InterestNetworkViz.from_interest_analysis(data, title=data.get('title', '利益关系网络'))
            self._figure_counter += 1
            chart_title = data.get('title', '利益关系网络')
            self._chart_entries.append((self._figure_counter, chart_title))
            viz.embed_to_word(doc, width_inches=5.5, figure_number=self._figure_counter)

            # 同时保存交互式 HTML 到报告目录
            out_dir = report_dir or self.export_dir
            safe_title = re.sub(r'[\\/*?:"<>|]', '_', data.get('title', 'network'))[:30]
            html_path = os.path.join(out_dir, f"{safe_title}_交互式.html")
            viz.save_html(html_path, animated=True)

            return True
        except Exception as e:
            # 解析失败，静默回退到字符画
            return False

    # ---------- PDF 导出（v4.0 六部分结构） ----------

    # ========== 从 Word 转 PDF（使用 Word COM，格式完全保真） ==========

    def _add_hyperlink(self, paragraph, text, url, font_name='仿宋', font_size=None, color=None):
        """在 Word 段落中添加可点击的超链接"""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True
        )

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # 颜色（RGBColor 是元组，用下标访问）
        if color:
            c = OxmlElement('w:color')
            r_val = '{:02X}{:02X}{:02X}'.format(color[0], color[1], color[2])
            c.set(qn('w:val'), r_val)
            rPr.append(c)

        # 下划线（超链接视觉提示）
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        # 中文字体
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)

        # 字号
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)

        new_run.append(rPr)

        t = OxmlElement('w:t')
        t.text = text
        new_run.append(t)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _render_reference_line(self, paragraph, line, font_size=None, color=None):
        """
        解析参考文献的一行文本，自动识别超链接并渲染：
        - [显示文本](http://...)  → 显示文本是可点击超链接
        - 裸 URL http(s)://...  → URL 本身是可点击超链接
        - 其余部分 → 普通文本
        """
        line = line.strip()
        if not line:
            return

        # 先处理 Markdown 链接 [text](url)，再处理裸 URL
        # 用正则把 line 拆成片段：Markdown 链接 / 裸 URL / 普通文本
        md_pat = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
        url_pat = re.compile(r'(https?://[^\s\[\]()]+)')

        # 收集所有特殊片段的位置，按出现顺序合并
        segments = []
        # 优先找 Markdown 链接
        for m in md_pat.finditer(line):
            segments.append({'start': m.start(), 'end': m.end(),
                            'type': 'md', 'text': m.group(1), 'url': m.group(2)})
        # 找裸 URL，排除已在 Markdown 链接内的
        for m in url_pat.finditer(line):
            if any(s['start'] <= m.start() < s['end'] for s in segments):
                continue
            segments.append({'start': m.start(), 'end': m.end(),
                            'type': 'url', 'text': m.group(1), 'url': m.group(1)})
        segments.sort(key=lambda x: x['start'])

        pos = 0
        for seg in segments:
            # 输出 URL 前面的普通文本
            if pos < seg['start']:
                run = paragraph.add_run(line[pos:seg['start']])
                set_cn_font(run, '仿宋', size=font_size, color=color)
            # 输出超链接
            self._add_hyperlink(paragraph, seg['text'], seg['url'],
                                font_name='仿宋', font_size=font_size, color=color)
            pos = seg['end']

        # 输出剩余普通文本
        if pos < len(line):
            run = paragraph.add_run(line[pos:])
            set_cn_font(run, '仿宋', size=font_size, color=color)

    def _convert_word_to_pdf(self, word_path):
        """
        使用 Word COM 接口（win32com）将 Word .docx 直接转为 PDF。
        格式完全保真（利用 Word 自身渲染引擎）。
        需要本地安装 Microsoft Word。

        v2 增强：嵌入 TrueType 字体 + 缺失字体位图化 + PDF/A 兼容。
        """
        if not _HAS_PYWIN32:
            print('⚠️ pywin32 不可用，无法调用 Word COM 转换 PDF')
            return None

        word_path = os.path.abspath(word_path)
        if not os.path.exists(word_path):
            print(f'⚠️ Word 文件不存在: {word_path}')
            return None

        pdf_path = word_path.replace('.docx', '.pdf')

        try:
            word = win32.Dispatch('Word.Application')
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(word_path)

            # 嵌入字体：确保 PDF 在所有设备上正确显示中文
            try:
                doc.EmbedTrueTypeFonts = True
            except Exception:
                pass  # 某些 Word 版本不支持

            # wdExportFormatPDF = 17
            doc.ExportAsFixedFormat(
                OutputFileName=pdf_path,
                ExportFormat=17,
                OptimizeFor=0,           # wdExportOptimizeForPrint
                BitmapMissingFonts=True, # 缺失字体位图化，防止乱码
                UseISO19005_1=False,     # 非 PDF/A（允许嵌入所有字体）
                DocStructureTags=True,   # 保留文档结构
            )
            doc.Close(SaveChanges=False)
            word.Quit()
            print(f'PDF 已生成（Word 转档，字体已嵌入）: {pdf_path}')
            return pdf_path
        except Exception as e:
            print(f'⚠️ Word 转 PDF 失败: {e}')
            return None

    # ============================================================
    #  利益关系网络可视化快捷方法
    # ============================================================

    def generate_interest_network(self, nodes, edges, title="利益关系网络", output_dir=None):
        """
        生成利益关系网络可视化图表（PNG + 交互式 HTML）。

        Args:
            nodes: list of dict, 每个节点 {"id": "A", "label": "...", "type": "political", "detail": "..."}
            edges: list of dict, 每条边 {"source": "A", "target": "B", "label": "...", "type": "economic"}
            title: 图表标题
            output_dir: 输出目录，默认使用 self.export_dir

        Returns:
            dict: {"png": "路径", "html": "路径", "base64": "data:image/png;base64,..."}
        """
        if not _HAS_VIZ:
            print('⚠️ interest_network_viz 模块不可用，无法生成网络可视化')
            return None

        viz = InterestNetworkViz.from_interest_analysis(
            {"nodes": nodes, "edges": edges},
            title=title
        )

        target = output_dir or self.export_dir
        os.makedirs(target, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = re.sub(r'[\\/*?:"<>|]', '_', title)[:30]
        prefix = f"{safe_title}_{timestamp}"

        png_path = os.path.join(target, f"{prefix}.png")
        html_path = os.path.join(target, f"{prefix}.html")

        viz.save_png(png_path, width=1440, height=1080)
        viz.save_html(html_path, animated=True)
        b64 = viz.generate_base64_png(width=900, height=675)

        return {
            "png": png_path,
            "html": html_path,
            "base64": b64,
            "summary": viz.summary()
        }

    def get_viz_types_reference(self):
        """返回利益类型和动线类型的参考信息，供分析提示使用"""
        if not _HAS_VIZ:
            return {}
        return {
            "interest_types": {k: {"name": v["name"], "color": v["color"], "short_label": v["short_label"]}
                              for k, v in INTEREST_TYPES.items()},
            "flow_types": {k: {"name": v["name"], "color": v["color"]} for k, v in FLOW_TYPES.items()},
        }

    # ============================================================
    #  网络图片搜索（v6.1 新增，零成本版权合规）
    # ============================================================

    def search_images(self, query, count=3, report_dir=None):
        """
        链式降级搜索版权合规图片。

        Wikimedia Commons → Unsplash → 官媒配图（预留）

        Args:
            query: 搜索关键词（中英均可）
            count: 最多返回图片数量
            report_dir: 报告目录（决定 images/ 子目录位置）

        Returns:
            list[dict]: 每项 {local_path, source, license, source_url, attribution}
            空列表表示未找到合规图片。
        """
        if not _HAS_IMAGE_SOURCER:
            print("  [图片搜索] image_sourcing 模块不可用")
            return []

        target_dir = report_dir or self.export_dir
        sourcer = ImageSourcer(target_dir)
        results = sourcer.search(query, limit=count)

        if not results:
            print(f"  [图片搜索] 「{query}」→ 未找到合规图片（已尝试 Wikimedia → Unsplash）")
        else:
            print(f"  [图片搜索] 「{query}」→ 找到 {len(results)} 张合规图片")
            for r in results:
                print(f"    → {r['source']} | {r['license']} | {r.get('local_path','')}")

        return results


if __name__ == "__main__":
    engine = CaseAnalysisEngine("theory_config.json")
