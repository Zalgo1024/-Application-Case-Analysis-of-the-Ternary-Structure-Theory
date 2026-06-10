# 三元结构理论 案例分析引擎 — 利益网络可视化引擎
# Copyright (C) 2026 李政恒 (Li Zhengheng)
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU Affero General Public License as published
# by the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU Affero General Public License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

"""
利益关系网络可视化系统 v1.0
========================================
基于三元结构理论的利益维度，使用 SVG 图标/Base64 编码图片渲染网络元素。

核心特性：
- 六大利益类型节点（物质/安全/政治/身份文化/制度性未来/公共利益），各有独立 SVG 图标和配色
- 四种利益动线（经济/权力/文化/法律路径），各有独立线型和箭头样式
- 动态流向动画（CSS/SVG 动画）
- 交互式 HTML 输出（缩放/拖拽/点击查看详情）
- PNG 静态图输出（用于 Word/PDF 嵌入）
- Base64 编码输出

依赖：lxml, Pillow, python-docx（可选）
"""

import os
import re
import json
import math
import base64
import hashlib
from io import BytesIO
from datetime import datetime
from copy import deepcopy

try:
    from lxml import etree
    HAS_LXML = True
except ImportError:
    HAS_LXML = False

try:
    from PIL import Image, ImageDraw, ImageFont, ImageFilter
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

try:
    from docx.shared import Inches, Cm, Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============================================================
# 配色与图标定义
# ============================================================

INTEREST_TYPES = {
    "material": {
        "name": "物质利益",
        "color": "#D4A017",       # 金色
        "bg_color": "#FFF8DC",    # 玉米丝色背景
        "border_color": "#B8860B",
        "icon": "hexagon",        # 六边形 → 财富
        "short_label": "物",
    },
    "security": {
        "name": "安全利益",
        "color": "#2E8B57",       # 海绿色
        "bg_color": "#F0FFF0",    # 蜜瓜色背景
        "border_color": "#228B22",
        "icon": "shield",         # 盾牌
        "short_label": "安",
    },
    "political": {
        "name": "政治利益",
        "color": "#8B4586",       # 紫色
        "bg_color": "#FFF0F5",    # 薰衣草色背景
        "border_color": "#6A0DAD",
        "icon": "balance",        # 天平
        "short_label": "政",
    },
    "identity_culture": {
        "name": "身份与文化利益",
        "color": "#E07020",       # 橙色
        "bg_color": "#FFF5EE",    # 贝壳色背景
        "border_color": "#CC5500",
        "icon": "masks",          # 面具
        "short_label": "文",
    },
    "institutional_future": {
        "name": "制度性与未来利益",
        "color": "#3A6EA5",       # 钢蓝色
        "bg_color": "#F0F8FF",    # 爱丽丝蓝背景
        "border_color": "#1B4F72",
        "icon": "pillar",         # 柱子
        "short_label": "制",
    },
    "public": {
        "name": "公共利益",
        "color": "#20B2AA",       # 浅海洋绿
        "bg_color": "#F0FFFF",    # 天蓝色背景
        "border_color": "#008B8B",
        "icon": "globe",          # 地球
        "short_label": "公",
    },
}

FLOW_TYPES = {
    "economic": {
        "name": "经济路径",
        "color": "#D4A017",
        "dash": "none",           # 实线
        "width": 2.5,
        "arrow_size": 10,
        "label_prefix": "$",
    },
    "power": {
        "name": "权力路径",
        "color": "#8B4586",
        "dash": "8,4",            # 长虚线
        "width": 2.5,
        "arrow_size": 10,
        "label_prefix": "P",
    },
    "cultural": {
        "name": "文化路径",
        "color": "#E07020",
        "dash": "4,3",            # 短虚线
        "width": 2.0,
        "arrow_size": 8,
        "label_prefix": "C",
    },
    "legal": {
        "name": "法律路径",
        "color": "#3A6EA5",
        "dash": "2,2,6,2",        # 点划线
        "width": 2.0,
        "arrow_size": 8,
        "label_prefix": "L",
    },
}


# ============================================================
# SVG 图标路径生成器
# ============================================================

def _svg_icon_hexagon(cx, cy, r=18):
    """六边形图标（物质利益）"""
    pts = []
    for i in range(6):
        angle = math.pi / 6 + i * math.pi / 3
        pts.append(f"{cx + r * math.cos(angle):.1f},{cy + r * math.sin(angle):.1f}")
    return f'<polygon points="{" ".join(pts)}" />'

def _svg_icon_shield(cx, cy, r=18):
    """盾牌图标（安全利益）"""
    w, h = r * 0.9, r * 1.1
    path = (
        f"M {cx:.1f},{cy - h:.1f} "
        f"L {cx + w:.1f},{cy - h * 0.4:.1f} "
        f"L {cx + w * 0.7:.1f},{cy + h * 0.5:.1f} "
        f"Q {cx:.1f},{cy + h * 1.1:.1f} {cx - w * 0.7:.1f},{cy + h * 0.5:.1f} "
        f"L {cx - w:.1f},{cy - h * 0.4:.1f} Z"
    )
    return f'<path d="{path}" />'

def _svg_icon_balance(cx, cy, r=18):
    """天平图标（政治利益）"""
    w = r * 0.9
    lines = [
        # 中柱
        f'<line x1="{cx:.1f}" y1="{cy - r:.1f}" x2="{cx:.1f}" y2="{cy + r * 0.6:.1f}" stroke="white" stroke-width="2.5"/>',
        # 底座
        f'<line x1="{cx - w * 0.6:.1f}" y1="{cy + r * 0.6:.1f}" x2="{cx + w * 0.6:.1f}" y2="{cy + r * 0.6:.1f}" stroke="white" stroke-width="2.5"/>',
        # 横梁
        f'<line x1="{cx - w:.1f}" y1="{cy - r * 0.4:.1f}" x2="{cx + w:.1f}" y2="{cy - r * 0.6:.1f}" stroke="white" stroke-width="2"/>',
        # 左盘
        f'<path d="M {cx - w * 1.1:.1f},{cy - r * 0.2:.1f} Q {cx - w * 0.6:.1f},{cy + r * 0.1:.1f} {cx - w * 0.1:.1f},{cy - r * 0.2:.1f}" fill="none" stroke="white" stroke-width="1.5"/>',
        # 右盘
        f'<path d="M {cx + w * 0.1:.1f},{cy - r * 0.4:.1f} Q {cx + w * 0.6:.1f},{cy - r * 0.1:.1f} {cx + w * 1.1:.1f},{cy - r * 0.4:.1f}" fill="none" stroke="white" stroke-width="1.5"/>',
    ]
    return '\n    '.join(lines)

def _svg_icon_masks(cx, cy, r=18):
    """双面具图标（身份与文化利益）"""
    r1 = r * 0.55
    r2 = r * 0.55
    offset = r * 0.4
    return (
        f'<circle cx="{cx - offset:.1f}" cy="{cy:.1f}" r="{r1:.1f}" fill="none" stroke="white" stroke-width="2"/>'
        f'<circle cx="{cx + offset:.1f}" cy="{cy:.1f}" r="{r2:.1f}" fill="none" stroke="white" stroke-width="2"/>'
        f'<line x1="{cx - offset - r1 * 0.3:.1f}" y1="{cy - r1 * 0.2:.1f}" x2="{cx - offset + r1 * 0.3:.1f}" y2="{cy - r1 * 0.2:.1f}" stroke="white" stroke-width="1.5"/>'
        f'<line x1="{cx + offset - r2 * 0.3:.1f}" y1="{cy + r2 * 0.15:.1f}" x2="{cx + offset + r2 * 0.3:.1f}" y2="{cy + r2 * 0.15:.1f}" stroke="white" stroke-width="1.5"/>'
    )

def _svg_icon_pillar(cx, cy, r=18):
    """柱子图标（制度性与未来利益）"""
    w = r * 0.5
    return (
        f'<rect x="{cx - w:.1f}" y="{cy - r * 0.9:.1f}" width="{w * 2:.1f}" height="{r * 1.5:.1f}" rx="2" fill="none" stroke="white" stroke-width="2"/>'
        f'<line x1="{cx - w * 1.3:.1f}" y1="{cy - r * 0.9:.1f}" x2="{cx + w * 1.3:.1f}" y2="{cy - r * 0.9:.1f}" stroke="white" stroke-width="2.5"/>'
        f'<line x1="{cx - w * 1.3:.1f}" y1="{cy + r * 0.6:.1f}" x2="{cx + w * 1.3:.1f}" y2="{cy + r * 0.6:.1f}" stroke="white" stroke-width="2.5"/>'
    )

def _svg_icon_globe(cx, cy, r=18):
    """地球图标（公共利益）"""
    return (
        f'<circle cx="{cx:.1f}" cy="{cy:.1f}" r="{r * 0.85:.1f}" fill="none" stroke="white" stroke-width="2"/>'
        f'<ellipse cx="{cx:.1f}" cy="{cy:.1f}" rx="{r * 0.4:.1f}" ry="{r * 0.85:.1f}" fill="none" stroke="white" stroke-width="1.2"/>'
        f'<line x1="{cx - r * 0.85:.1f}" y1="{cy:.1f}" x2="{cx + r * 0.85:.1f}" y2="{cy:.1f}" stroke="white" stroke-width="1.2"/>'
        f'<path d="M {cx - r * 0.75:.1f},{cy - r * 0.45:.1f} Q {cx:.1f},{cy - r * 0.3:.1f} {cx + r * 0.75:.1f},{cy - r * 0.45:.1f}" fill="none" stroke="white" stroke-width="1"/>'
        f'<path d="M {cx - r * 0.75:.1f},{cy + r * 0.45:.1f} Q {cx:.1f},{cy + r * 0.3:.1f} {cx + r * 0.75:.1f},{cy + r * 0.45:.1f}" fill="none" stroke="white" stroke-width="1"/>'
    )

ICON_GENERATORS = {
    "hexagon": _svg_icon_hexagon,
    "shield": _svg_icon_shield,
    "balance": _svg_icon_balance,
    "masks": _svg_icon_masks,
    "pillar": _svg_icon_pillar,
    "globe": _svg_icon_globe,
}


# ============================================================
# 布局算法
# ============================================================

def _layout_circular(nodes, width, height, margin=80):
    """圆形布局：所有节点均匀分布在一个圆上"""
    cx, cy = width / 2, height / 2
    r = min(width, height) / 2 - margin
    n = len(nodes)
    if n == 0:
        return {}
    positions = {}
    for i, node in enumerate(nodes):
        angle = 2 * math.pi * i / n - math.pi / 2
        x = cx + r * math.cos(angle)
        y = cy + r * math.sin(angle)
        positions[node['id']] = (x, y)
    return positions


def _layout_force(nodes, edges, width, height, margin=80, iterations=200):
    """
    力导向布局（v2 增强版）。
    节点间有斥力，边连节点间有引力，所有节点被限制在画布内。
    v2 改进：
    - 增加迭代次数至 200，提高收敛质量
    - 增大斥力系数，减少节点重叠
    - 添加重叠检测与修正（节点间距过小时强制推开）
    - 引入居中引力，防止节点漂移到边缘
    - 优化退火策略，前期大幅移动，后期精细调整
    """
    cx, cy = width / 2, height / 2
    n = len(nodes)
    if n == 0:
        return {}

    # 初始位置：圆形散布
    positions = {}
    for i, node in enumerate(nodes):
        angle = 2 * math.pi * i / n - math.pi / 2
        r0 = min(width, height) * 0.3
        positions[node['id']] = (cx + r0 * math.cos(angle), cy + r0 * math.sin(angle))

    if n <= 1:
        return positions

    # 构建邻接表
    adj = {node['id']: set() for node in nodes}
    for edge in edges:
        src, tgt = edge['source'], edge['target']
        if src in adj and tgt in adj:
            adj[src].add(tgt)
            adj[tgt].add(src)

    # 最小节点间距（避免重叠，基于节点标签长度和固定半径估算）
    min_dist = 90  # 节点间最小间距（像素）

    k = math.sqrt((width - 2 * margin) * (height - 2 * margin) / max(n, 1)) * 1.0

    for iteration in range(iterations):
        # 改进的退火策略：前期大幅移动，后期精细调整
        progress = iteration / iterations
        if progress < 0.3:
            temp = k * 0.8
        elif progress < 0.7:
            temp = k * 0.4
        else:
            temp = max(k * 0.05, k * 0.1 * (1.0 - progress))

        disp = {nid: [0.0, 0.0] for nid in positions}

        # 斥力（增大系数）
        ids = list(positions.keys())
        for i in range(len(ids)):
            for j in range(i + 1, len(ids)):
                ni, nj = ids[i], ids[j]
                dx = positions[ni][0] - positions[nj][0]
                dy = positions[ni][1] - positions[nj][1]
                dist = max(math.sqrt(dx * dx + dy * dy), 1.0)

                # 标准斥力
                force = k * k / dist

                # 重叠惩罚：距离小于 min_dist 时额外施加斥力
                if dist < min_dist:
                    overlap_force = (min_dist - dist) * 2.0
                    force += overlap_force

                fx = dx / dist * force
                fy = dy / dist * force
                disp[ni][0] += fx
                disp[ni][1] += fy
                disp[nj][0] -= fx
                disp[nj][1] -= fy

        # 引力（边连节点）
        for edge in edges:
            src, tgt = edge['source'], edge['target']
            if src not in positions or tgt not in positions:
                continue
            dx = positions[src][0] - positions[tgt][0]
            dy = positions[src][1] - positions[tgt][1]
            dist = max(math.sqrt(dx * dx + dy * dy), 1.0)
            force = dist * dist / k
            fx = dx / dist * force
            fy = dy / dist * force
            disp[src][0] -= fx * 0.3
            disp[src][1] -= fy * 0.3
            disp[tgt][0] += fx * 0.3
            disp[tgt][1] += fy * 0.3

        # 居中引力（防止节点漂移到边缘）
        for nid in positions:
            dx = cx - positions[nid][0]
            dy = cy - positions[nid][1]
            dist = max(math.sqrt(dx * dx + dy * dy), 1.0)
            gravity = dist * 0.01  # 微弱的向心力
            disp[nid][0] += dx / dist * gravity
            disp[nid][1] += dy / dist * gravity

        # 应用位移
        for nid in positions:
            dx, dy = disp[nid]
            dist = max(math.sqrt(dx * dx + dy * dy), 0.01)
            x = positions[nid][0] + dx / dist * min(dist, temp)
            y = positions[nid][1] + dy / dist * min(dist, temp)
            x = max(margin, min(width - margin, x))
            y = max(margin, min(height - margin, y))
            positions[nid] = (x, y)

    return positions


def _layout_hierarchical(nodes, edges, width, height, margin=80):
    """层级布局：按拓扑排序分层"""
    n = len(nodes)
    if n == 0:
        return {}

    node_ids = [nd['id'] for nd in nodes]
    in_degree = {nid: 0 for nid in node_ids}
    children = {nid: [] for nid in node_ids}
    for edge in edges:
        src, tgt = edge['source'], edge['target']
        if src in in_degree and tgt in in_degree:
            in_degree[tgt] += 1
            children[src].append(tgt)

    # BFS 分层
    layers = []
    visited = set()
    queue = [nid for nid in node_ids if in_degree[nid] == 0]
    if not queue:
        queue = [node_ids[0]]

    while queue:
        layer = []
        next_queue = []
        for nid in queue:
            if nid in visited:
                continue
            visited.add(nid)
            layer.append(nid)
            for child in children[nid]:
                if child not in visited:
                    in_degree[child] -= 1
                    if in_degree[child] <= 0:
                        next_queue.append(child)
        if layer:
            layers.append(layer)
        queue = next_queue

    # 未访问的节点放入最后一层
    for nid in node_ids:
        if nid not in visited:
            if layers:
                layers[-1].append(nid)
            else:
                layers.append([nid])

    num_layers = len(layers)
    positions = {}
    usable_h = height - 2 * margin
    usable_w = width - 2 * margin

    for li, layer in enumerate(layers):
        y = margin + (usable_h * (li + 0.5) / num_layers) if num_layers > 0 else height / 2
        for ni, nid in enumerate(layer):
            x = margin + usable_w * (ni + 0.5) / max(len(layer), 1)
            positions[nid] = (x, y)

    return positions


# ============================================================
# 主类：InterestNetworkViz
# ============================================================

class InterestNetworkViz:
    """
    利益关系网络可视化器。

    用法：
        viz = InterestNetworkViz(title="案例A利益关系网络")
        viz.add_node("A", "地方政府", "political", detail="掌握审批权和土地分配")
        viz.add_node("B", "开发商", "material", detail="追求土地增值利润")
        viz.add_edge("A", "B", "土地审批→利润分成", "economic")
        viz.add_edge("B", "A", "回扣/税收贡献", "power")
        html = viz.generate_html()  # 交互式HTML
        viz.save_html("network.html")
        viz.embed_to_word(doc)     # 嵌入Word文档
    """

    NODE_RADIUS = 28
    LEGEND_HEIGHT = 160  # 底部图例区独立高度（不参与力导向布局）

    def __init__(self, title="利益关系网络", width=900, height=650, layout="force"):
        self.title = title
        self.width = width
        self.height = height
        self.layout_algo = layout
        # 总画布高度 = 布局区 + 图例区（图例从不覆盖节点）
        self.total_height = self.height + self.LEGEND_HEIGHT
        self.nodes = []
        self.edges = []
        self._node_index = {}  # id -> node dict

    # ---------- 数据构建 ----------

    def add_node(self, node_id, label, interest_type, detail="", group=""):
        """
        添加利益节点。

        Args:
            node_id: 节点唯一标识
            label: 显示标签（如"地方政府"）
            interest_type: 利益类型（material/security/political/identity_culture/institutional_future/public）
            detail: 点击查看的详细信息
            group: 可选分组标签
        """
        if interest_type not in INTEREST_TYPES:
            interest_type = "material"  # 回退
        node = {
            "id": node_id,
            "label": label,
            "type": interest_type,
            "detail": detail or label,
            "group": group,
        }
        self.nodes.append(node)
        self._node_index[node_id] = node
        return self

    def add_edge(self, source, target, label="", flow_type="economic", detail="", bidirectional=False):
        """
        添加利益动线（边）。

        Args:
            source: 源节点ID
            target: 目标节点ID
            label: 动线标签
            flow_type: 动线类型（economic/power/cultural/legal）
            detail: 详细说明
            bidirectional: 是否双向
        """
        if flow_type not in FLOW_TYPES:
            flow_type = "economic"
        edge = {
            "source": source,
            "target": target,
            "label": label,
            "type": flow_type,
            "detail": detail or label,
            "bidirectional": bidirectional,
        }
        self.edges.append(edge)
        return self

    # ---------- 布局 ----------

    def _compute_positions(self):
        """根据布局算法计算节点位置（图例区不参与布局，永不重叠）"""
        if self.layout_algo == "circular":
            return _layout_circular(self.nodes, self.width, self.height)
        elif self.layout_algo == "hierarchical":
            return _layout_hierarchical(self.nodes, self.edges, self.width, self.height)
        else:
            return _layout_force(self.nodes, self.edges, self.width, self.height)

    # ---------- SVG 生成 ----------

    def generate_svg(self, animated=True, embed_style=True):
        """
        生成完整的 SVG 字符串。

        Args:
            animated: 是否包含流向动画
            embed_style: 是否内嵌样式
        """
        positions = self._compute_positions()
        type_config = INTEREST_TYPES
        flow_config = FLOW_TYPES
        R = self.NODE_RADIUS

        total_h = self.total_height
        svg_parts = []

        # SVG 头（viewBox 包含底部图例区）
        svg_parts.append(
            f'<svg xmlns="http://www.w3.org/2000/svg" '
            f'viewBox="0 0 {self.width} {total_h}" '
            f'width="{self.width}" height="{total_h}" '
            f'style="font-family: \'Microsoft YaHei\', \'SimHei\', \'PingFang SC\', sans-serif;">'
        )

        # 内嵌样式
        if embed_style:
            svg_parts.append(self._svg_styles(animated))

        # 背景（覆盖全画布，包括图例区）
        svg_parts.append(
            f'<rect width="{self.width}" height="{total_h}" fill="#FAFBFC" rx="8"/>'
        )
        # 标题
        svg_parts.append(
            f'<text x="{self.width / 2}" y="32" text-anchor="middle" '
            f'font-size="22" font-weight="bold" fill="#1B3A5C">{self.title}</text>'
        )

        # Defs: 箭头标记
        svg_parts.append('<defs>')
        for ft_key, ft in flow_config.items():
            # 正向箭头
            svg_parts.append(
                f'<marker id="arrow-{ft_key}" viewBox="0 0 {ft["arrow_size"]} {ft["arrow_size"]}" '
                f'refX="{ft["arrow_size"]}" refY="{ft["arrow_size"] / 2}" '
                f'markerWidth="{ft["arrow_size"]}" markerHeight="{ft["arrow_size"]}" orient="auto">'
                f'<path d="M0,0 L{ft["arrow_size"]},{ft["arrow_size"] / 2} L0,{ft["arrow_size"]}" '
                f'fill="{ft["color"]}"/></marker>'
            )
            # 反向箭头（双向边用）
            svg_parts.append(
                f'<marker id="arrow-rev-{ft_key}" viewBox="0 0 {ft["arrow_size"]} {ft["arrow_size"]}" '
                f'refX="0" refY="{ft["arrow_size"] / 2}" '
                f'markerWidth="{ft["arrow_size"]}" markerHeight="{ft["arrow_size"]}" orient="auto">'
                f'<path d="M{ft["arrow_size"]},0 L0,{ft["arrow_size"] / 2} L{ft["arrow_size"]},{ft["arrow_size"]}" '
                f'fill="{ft["color"]}"/></marker>'
            )
        svg_parts.append('</defs>')

        # 绘制边（利益动线）
        for edge in self.edges:
            src_pos = positions.get(edge['source'])
            tgt_pos = positions.get(edge['target'])
            if not src_pos or not tgt_pos:
                continue
            ft = flow_config[edge['type']]
            sx, sy = src_pos
            tx, ty = tgt_pos

            # 计算边的起止点（从节点边缘开始，不穿过节点）
            dx = tx - sx
            dy = ty - sy
            dist = max(math.sqrt(dx * dx + dy * dy), 1)
            ux, uy = dx / dist, dy / dist

            # 如果源和目标是同一个，画自环
            if edge['source'] == edge['target']:
                loop_r = R * 1.2
                path_d = (
                    f"M {sx + R * 0.7:.1f},{sy - R * 0.5:.1f} "
                    f"C {sx + R * 2:.1f},{sy - R * 2:.1f} "
                    f"{sx - R * 2:.1f},{sy - R * 2:.1f} "
                    f"{sx - R * 0.7:.1f},{sy - R * 0.5:.1f}"
                )
                svg_parts.append(
                    f'<path d="{path_d}" fill="none" stroke="{ft["color"]}" '
                    f'stroke-width="{ft["width"]}" stroke-dasharray="{ft["dash"]}" '
                    f'marker-end="url(#arrow-{edge["type"]})"/>'
                )
                # 标签
                svg_parts.append(
                    f'<text x="{sx:.1f}" y="{sy - R * 2.2:.1f}" text-anchor="middle" '
                    f'font-size="14" fill="{ft["color"]}">{edge["label"]}</text>'
                )
                continue

            # 检查是否有同向边（多条边需要偏移）
            offset_x, offset_y = 0, 0
            same_dir = sum(1 for e in self.edges
                          if (e['source'] == edge['source'] and e['target'] == edge['target'])
                          or (e['source'] == edge['target'] and e['target'] == edge['source']))
            if same_dir > 1:
                # 垂直偏移
                perp_x, perp_y = -uy, ux
                idx = sum(1 for e in self.edges[:self.edges.index(edge)]
                         if (e['source'] == edge['source'] and e['target'] == edge['target'])
                         or (e['source'] == edge['target'] and e['target'] == edge['source']))
                offset_x = perp_x * 15 * (idx - same_dir / 2 + 0.5)
                offset_y = perp_y * 15 * (idx - same_dir / 2 + 0.5)

            # 边起止点
            start_x = sx + ux * R + offset_x
            start_y = sy + uy * R + offset_y
            end_x = tx - ux * R + offset_x
            end_y = ty - uy * R + offset_y

            # 控制点（曲线）
            mid_x = (start_x + end_x) / 2 + offset_x * 0.5
            mid_y = (start_y + end_y) / 2 + offset_y * 0.5
            # 轻微弯曲
            curve_offset = 20 if same_dir > 1 else 8
            ctrl_x = mid_x + (-uy) * curve_offset
            ctrl_y = mid_y + ux * curve_offset

            path_d = (
                f"M {start_x:.1f},{start_y:.1f} "
                f"Q {ctrl_x:.1f},{ctrl_y:.1f} "
                f"{end_x:.1f},{end_y:.1f}"
            )

            svg_parts.append(
                f'<path class="flow-path flow-{edge["type"]}" d="{path_d}" '
                f'fill="none" stroke="{ft["color"]}" '
                f'stroke-width="{ft["width"]}" stroke-dasharray="{ft["dash"]}" '
                f'marker-end="url(#arrow-{edge["type"]})"/>'
            )

            # 双向箭头
            if edge.get('bidirectional'):
                svg_parts.append(
                    f'<path d="{path_d}" fill="none" stroke="{ft["color"]}" '
                    f'stroke-width="{ft["width"]}" stroke-dasharray="{ft["dash"]}" '
                    f'marker-start="url(#arrow-rev-{edge["type"]})" opacity="0"/>'
                )

            # 动线标签
            if edge['label']:
                label_x = (start_x + end_x) / 2 + (-uy) * (curve_offset + 12)
                label_y = (start_y + end_y) / 2 + ux * (curve_offset + 12)
                svg_parts.append(
                    f'<text x="{label_x:.1f}" y="{label_y:.1f}" text-anchor="middle" '
                    f'font-size="14" fill="{ft["color"]}" font-weight="500">'
                    f'{edge["label"]}</text>'
                )

        # 绘制节点
        for node in self.nodes:
            pos = positions.get(node['id'])
            if not pos:
                continue
            nx, ny = pos
            tc = type_config[node['type']]
            R_node = R

            # 节点组（可交互）
            svg_parts.append(
                f'<g class="node-group" data-id="{node["id"]}" '
                f'data-type="{node["type"]}" data-detail="{self._escape_attr(node["detail"])}">'
            )

            # 光晕（阴影效果）
            svg_parts.append(
                f'<circle cx="{nx:.1f}" cy="{ny:.1f}" r="{R_node + 4:.1f}" '
                f'fill="{tc["color"]}" opacity="0.15"/>'
            )

            # 背景圆
            svg_parts.append(
                f'<circle cx="{nx:.1f}" cy="{ny:.1f}" r="{R_node:.1f}" '
                f'fill="{tc["bg_color"]}" stroke="{tc["border_color"]}" stroke-width="2.5"/>'
            )

            # 图标
            icon_gen = ICON_GENERATORS.get(tc['icon'])
            if icon_gen:
                icon_svg = icon_gen(nx, ny, R_node * 0.55)
                if icon_svg.strip().startswith('<'):
                    svg_parts.append(icon_svg)
                else:
                    svg_parts.append(
                        f'<g fill="{tc["color"]}" stroke="white" stroke-width="1.5">{icon_svg}</g>'
                    )

            # 类型标签（小字，节点内上方）
            svg_parts.append(
                f'<text x="{nx:.1f}" y="{ny - R_node - 6:.1f}" text-anchor="middle" '
                f'font-size="13" fill="{tc["border_color"]}" font-weight="600">'
                f'{tc["short_label"]}</text>'
            )

            # 名称标签（节点外下方）
            svg_parts.append(
                f'<text x="{nx:.1f}" y="{ny + R_node + 16:.1f}" text-anchor="middle" '
                f'font-size="15" fill="#333" font-weight="500">'
                f'{node["label"]}</text>'
            )

            svg_parts.append('</g>')

        # 图例
        svg_parts.append(self._svg_legend(type_config, flow_config))

        svg_parts.append('</svg>')
        return '\n'.join(svg_parts)

    def _svg_styles(self, animated=True):
        """生成 SVG 内嵌 CSS 样式"""
        anim_css = ""
        if animated:
            for ft_key in FLOW_TYPES:
                anim_css += f"""
    .flow-{ft_key} {{
        stroke-dashoffset: 100;
        animation: flowAnim-{ft_key} 2s linear infinite;
    }}
    @keyframes flowAnim-{ft_key} {{
        to {{ stroke-dashoffset: 0; }}
    }}"""

        return f'''<defs><style>
    .node-group {{ cursor: pointer; }}
    .node-group:hover circle {{ filter: brightness(1.1); stroke-width: 3.5; }}
    .node-group:hover text {{ font-weight: 700; }}
    .flow-path {{ transition: stroke-width 0.2s; }}
    .flow-path:hover {{ stroke-width: 4; }}
    .legend-box {{ cursor: default; }}{anim_css}
</style></defs>'''

    def _svg_legend(self, type_config, flow_config):
        """生成图例 — 底部横向条带，与节点区域物理隔离，永不遮挡"""
        legend_y = self.height + 10  # 图例起始 y（中间有 10px 间距）
        lw = self.width - 30  # 图例可用宽度

        parts = ['<g class="legend-box">']

        # 分隔线（节点区 / 图例区）
        parts.append(
            f'<line x1="15" y1="{self.height}" x2="{self.width - 15}" y2="{self.height}" '
            f'stroke="#E0E0E0" stroke-width="1" stroke-dasharray="6,3"/>'
        )

        # --- 第一行：利益类型（6个水平排列） ---
        parts.append(
            f'<text x="15" y="{legend_y + 16}" font-size="14" '
            f'font-weight="bold" fill="#333">利益类型</text>'
        )
        type_items = list(type_config.items())
        for i, (tkey, tconf) in enumerate(type_items):
            # 6 个均匀分布在 label 右侧
            start_x = 100
            spacing = 155  # v6.1.1: 125→155 容纳完整中文名称
            cx = start_x + i * spacing
            cy = legend_y + 10
            parts.append(
                f'<circle cx="{cx}" cy="{cy}" r="6" fill="{tconf["color"]}" '
                f'stroke="{tconf["border_color"]}" stroke-width="1.5"/>'
            )
            parts.append(
                f'<text x="{cx + 10}" y="{cy + 4}" font-size="13" fill="#555">'
                f'{tconf["short_label"]} {tconf["name"]}</text>'
            )

        # --- 第二行：利益动线（4个水平排列） ---
        ly2 = legend_y + 38
        parts.append(
            f'<text x="15" y="{ly2 + 14}" font-size="14" '
            f'font-weight="bold" fill="#333">利益动线</text>'
        )
        flow_items = list(flow_config.items())
        for i, (fkey, fconf) in enumerate(flow_items):
            start_x = 100
            spacing = 190
            lx = start_x + i * spacing
            ly = ly2 + 8
            parts.append(
                f'<line x1="{lx}" y1="{ly}" x2="{lx + 35}" y2="{ly}" '
                f'stroke="{fconf["color"]}" stroke-width="2.5" '
                f'stroke-dasharray="{fconf["dash"]}"/>'
            )
            parts.append(
                f'<text x="{lx + 42}" y="{ly + 4}" font-size="13" fill="#555">'
                f'{fconf["name"]}</text>'
            )

        # --- 第三行：交互提示 ---
        ly3 = legend_y + 68
        parts.append(
            f'<text x="{self.width - 20}" y="{ly3 + 12}" text-anchor="end" '
            f'font-size="11" fill="#999">拖拽节点 · 滚轮缩放 · 点击查看详情</text>'
        )

        parts.append('</g>')
        return '\n'.join(parts)

    @staticmethod
    def _escape_attr(text):
        """转义 XML 属性中的特殊字符"""
        return text.replace('&', '&amp;').replace('"', '&quot;').replace('<', '&lt;').replace('>', '&gt;')

    # ---------- HTML 交互页面生成 ----------

    def generate_html(self, animated=True):
        """
        生成完整的交互式 HTML 页面。
        支持：缩放、拖拽、点击查看详情。
        """
        svg_content = self.generate_svg(animated=animated, embed_style=False)
        # 移除 SVG 外层标签（嵌入 HTML 中时由 JS 控制）
        svg_inner = svg_content
        if svg_content.startswith('<svg'):
            svg_inner = svg_content[svg_content.find('>') + 1:]
            if svg_inner.endswith('</svg>'):
                svg_inner = svg_inner[:-6]

        html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{self.title} - 利益关系网络</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: 'Microsoft YaHei', 'SimHei', 'PingFang SC', sans-serif;
    background: #F0F2F5;
    color: #333;
    overflow: hidden;
}}
.header {{
    background: linear-gradient(135deg, #1B3A5C, #2C5F8A);
    color: white;
    padding: 12px 24px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    box-shadow: 0 2px 8px rgba(0,0,0,0.15);
}}
.header h1 {{ font-size: 18px; font-weight: 600; }}
.header .controls {{ display: flex; gap: 8px; }}
.header .controls button {{
    background: rgba(255,255,255,0.2);
    border: 1px solid rgba(255,255,255,0.3);
    color: white;
    padding: 4px 12px;
    border-radius: 4px;
    cursor: pointer;
    font-size: 13px;
    transition: background 0.2s;
}}
.header .controls button:hover {{ background: rgba(255,255,255,0.35); }}
.toolbar {{
    background: white;
    padding: 6px 16px;
    border-bottom: 1px solid #E0E0E0;
    display: flex;
    gap: 12px;
    align-items: center;
    font-size: 12px;
}}
.toolbar label {{ display: flex; align-items: center; gap: 4px; cursor: pointer; }}
.toolbar input[type="checkbox"] {{ cursor: pointer; }}
.toolbar select {{
    padding: 2px 8px;
    border: 1px solid #CCC;
    border-radius: 3px;
    font-size: 12px;
}}
#canvas-container {{
    width: 100%;
    height: calc(100vh - 80px);
    position: relative;
    overflow: hidden;
    cursor: grab;
}}
#canvas-container:active {{ cursor: grabbing; }}
#canvas-container svg {{
    position: absolute;
    top: 0; left: 0;
    transform-origin: 0 0;
}}
.node-group {{ cursor: pointer; }}
.node-group:hover circle {{ filter: brightness(1.1); stroke-width: 3.5; }}
.node-group:hover text {{ font-weight: 700; }}
.flow-path {{ transition: stroke-width 0.2s; }}
.flow-path:hover {{ stroke-width: 4; }}
.legend-box {{ cursor: default; pointer-events: none; }}
{" ".join(f'.flow-{fk} {{ stroke-dashoffset: 100; animation: flowAnim-{fk} 2s linear infinite; }} @keyframes flowAnim-{fk} {{ to {{ stroke-dashoffset: 0; }} }}' for fk in FLOW_TYPES) if animated else ""}

/* 详情弹窗 */
#detail-panel {{
    position: fixed;
    right: -360px;
    top: 80px;
    width: 340px;
    max-height: calc(100vh - 100px);
    background: white;
    box-shadow: -2px 0 12px rgba(0,0,0,0.15);
    border-radius: 8px 0 0 8px;
    transition: right 0.3s ease;
    overflow-y: auto;
    z-index: 100;
}}
#detail-panel.open {{ right: 0; }}
#detail-panel .panel-header {{
    padding: 16px;
    border-bottom: 1px solid #EEE;
    display: flex;
    justify-content: space-between;
    align-items: center;
}}
#detail-panel .panel-header h3 {{ font-size: 15px; }}
#detail-panel .panel-header .close-btn {{
    background: none; border: none; font-size: 20px; cursor: pointer; color: #999;
}}
#detail-panel .panel-body {{ padding: 16px; }}
#detail-panel .type-badge {{
    display: inline-block;
    padding: 2px 10px;
    border-radius: 12px;
    font-size: 11px;
    font-weight: 600;
    margin-bottom: 8px;
}}
#detail-panel .detail-text {{ font-size: 13px; line-height: 1.8; color: #555; }}
#detail-panel .edges-section {{ margin-top: 12px; }}
#detail-panel .edge-item {{
    padding: 8px 10px;
    margin: 4px 0;
    border-left: 3px solid #CCC;
    background: #FAFAFA;
    border-radius: 0 4px 4px 0;
    font-size: 12px;
}}

/* 缩放提示 */
#zoom-indicator {{
    position: fixed;
    bottom: 20px;
    right: 20px;
    background: rgba(0,0,0,0.6);
    color: white;
    padding: 4px 12px;
    border-radius: 12px;
    font-size: 12px;
    pointer-events: none;
    opacity: 0;
    transition: opacity 0.3s;
}}
#zoom-indicator.show {{ opacity: 1; }}

/* 导出按钮区 */
.export-bar {{
    position: fixed;
    bottom: 20px;
    left: 20px;
    display: flex;
    gap: 8px;
    z-index: 50;
}}
.export-bar button {{
    background: white;
    border: 1px solid #DDD;
    padding: 6px 14px;
    border-radius: 6px;
    cursor: pointer;
    font-size: 12px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.1);
    transition: all 0.2s;
}}
.export-bar button:hover {{ background: #F0F0F0; border-color: #BBB; }}
</style>
</head>
<body>

<div class="header">
    <h1>{self.title}</h1>
    <div class="controls">
        <button onclick="resetView()">重置视图</button>
        <button onclick="fitAll()">适应全部</button>
    </div>
</div>

<div class="toolbar">
    <label>布局：
        <select id="layout-select" onchange="changeLayout(this.value)">
            <option value="force"{' selected' if self.layout_algo == 'force' else ''}>力导向</option>
            <option value="circular"{' selected' if self.layout_algo == 'circular' else ''}>圆形</option>
            <option value="hierarchical"{' selected' if self.layout_algo == 'hierarchical' else ''}>层级</option>
        </select>
    </label>
    <label><input type="checkbox" id="anim-toggle" {'checked' if animated else ''} onchange="toggleAnimation(this.checked)"> 动画</label>
    <label><input type="checkbox" id="label-toggle" checked onchange="toggleLabels(this.checked)"> 标签</label>
    <span style="color:#999">滚轮缩放 | 拖拽平移 | 点击节点查看详情</span>
</div>

<div id="canvas-container">
    <svg xmlns="http://www.w3.org/2000/svg"
         viewBox="0 0 {self.width} {self.total_height}"
         width="{self.width}" height="{self.total_height}"
         id="main-svg">
        {svg_inner}
    </svg>
</div>

<div id="detail-panel">
    <div class="panel-header">
        <h3 id="detail-title">节点详情</h3>
        <button class="close-btn" onclick="closeDetail()">&times;</button>
    </div>
    <div class="panel-body" id="detail-body"></div>
</div>

<div id="zoom-indicator">100%</div>

<div class="export-bar">
    <button onclick="exportSVG()">导出 SVG</button>
    <button onclick="exportPNG()">导出 PNG</button>
</div>

<script>
// === 交互系统 ===
const container = document.getElementById('canvas-container');
const svg = document.getElementById('main-svg');
const detailPanel = document.getElementById('detail-panel');
const zoomIndicator = document.getElementById('zoom-indicator');

let scale = 1, panX = 0, panY = 0;
let isDragging = false, dragStartX = 0, dragStartY = 0;
let isPanning = false;

// 节点数据
const nodesData = {json.dumps({n['id']: n for n in self.nodes}, ensure_ascii=False)};
const edgesData = {json.dumps(self.edges, ensure_ascii=False)};
const typeConfig = {json.dumps(INTEREST_TYPES, ensure_ascii=False)};

// 缩放
container.addEventListener('wheel', (e) => {{
    e.preventDefault();
    const delta = e.deltaY > 0 ? 0.9 : 1.1;
    const rect = container.getBoundingClientRect();
    const mx = e.clientX - rect.left;
    const my = e.clientY - rect.top;

    const newScale = Math.max(0.2, Math.min(5, scale * delta));
    const ratio = newScale / scale;

    panX = mx - ratio * (mx - panX);
    panY = my - ratio * (my - panY);
    scale = newScale;

    updateTransform();
    showZoom();
}}, {{ passive: false }});

// 拖拽平移
container.addEventListener('mousedown', (e) => {{
    if (e.target.closest('.node-group')) return;
    isPanning = true;
    dragStartX = e.clientX - panX;
    dragStartY = e.clientY - panY;
    container.style.cursor = 'grabbing';
}});
document.addEventListener('mousemove', (e) => {{
    if (!isPanning) return;
    panX = e.clientX - dragStartX;
    panY = e.clientY - dragStartY;
    updateTransform();
}});
document.addEventListener('mouseup', () => {{
    isPanning = false;
    container.style.cursor = 'grab';
}});

function updateTransform() {{
    svg.style.transform = `translate(${{panX}}px, ${{panY}}px) scale(${{scale}})`;
}}

function showZoom() {{
    zoomIndicator.textContent = Math.round(scale * 100) + '%';
    zoomIndicator.classList.add('show');
    clearTimeout(zoomIndicator._timer);
    zoomIndicator._timer = setTimeout(() => zoomIndicator.classList.remove('show'), 1200);
}}

function resetView() {{
    scale = 1; panX = 0; panY = 0;
    updateTransform();
    showZoom();
}}

function fitAll() {{
    const rect = container.getBoundingClientRect();
    const sx = rect.width / {self.width};
    const sy = rect.height / {self.total_height};
    scale = Math.min(sx, sy) * 0.9;
    panX = (rect.width - {self.width} * scale) / 2;
    panY = (rect.height - {self.total_height} * scale) / 2;
    updateTransform();
    showZoom();
}}

// 点击节点查看详情
document.querySelectorAll('.node-group').forEach(g => {{
    g.addEventListener('click', (e) => {{
        e.stopPropagation();
        const id = g.dataset.id;
        const type = g.dataset.type;
        const detail = g.dataset.detail;
        showDetail(id, type, detail);
    }});
}});

function showDetail(id, type, detail) {{
    const tc = typeConfig[type] || {{}};
    document.getElementById('detail-title').textContent = nodesData[id]?.label || id;
    const body = document.getElementById('detail-body');
    let html = `<span class="type-badge" style="background:${{tc.bg_color || '#EEE'}};color:${{tc.border_color || '#333'}}">${{tc.name || type}}</span>`;
    html += `<div class="detail-text">${{detail || '暂无详细信息'}}</div>`;

    // 关联边
    const related = edgesData.filter(e => e.source === id || e.target === id);
    if (related.length > 0) {{
        html += '<div class="edges-section"><strong>关联利益动线：</strong>';
        related.forEach(e => {{
            const dir = e.source === id ? '→' : '←';
            const other = e.source === id ? nodesData[e.target]?.label || e.target : nodesData[e.source]?.label || e.source;
            const ftName = ({' | '.join(f'"{fk}": "{fv["name"]}"' for fk, fv in FLOW_TYPES.items())})[e.type] || e.type;
            html += `<div class="edge-item" style="border-left-color:${{({'| '.join(f'"{fk}": "{fv["color"]}"' for fk, fv in FLOW_TYPES.items())})[e.type] || '#CCC'}}">${{dir}} ${{other}}（${{ftName}}）${{e.label ? '：' + e.label : ''}}</div>`;
        }});
        html += '</div>';
    }}
    body.innerHTML = html;
    detailPanel.classList.add('open');
}}

function closeDetail() {{
    detailPanel.classList.remove('open');
}}

// 动画控制
function toggleAnimation(on) {{
    document.querySelectorAll('.flow-path').forEach(p => {{
        p.style.animationPlayState = on ? 'running' : 'paused';
    }});
}}

// 标签控制
function toggleLabels(show) {{
    document.querySelectorAll('.node-group text, .flow-path + text').forEach(t => {{
        t.style.opacity = show ? 1 : 0;
    }});
}}

// 布局切换（简化：刷新页面）
function changeLayout(layout) {{
    const url = new URL(window.location);
    url.searchParams.set('layout', layout);
    window.location = url;
}}

// 导出
function exportSVG() {{
    const svgData = new XMLSerializer().serializeToString(svg);
    const blob = new Blob([svgData], {{ type: 'image/svg+xml' }});
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url; a.download = '{self.title}.svg'; a.click();
    URL.revokeObjectURL(url);
}}

function exportPNG() {{
    const svgData = new XMLSerializer().serializeToString(svg);
    const canvas = document.createElement('canvas');
    canvas.width = {self.width} * 2;
    canvas.height = {self.total_height} * 2;
    const ctx = canvas.getContext('2d');
    const img = new Image();
    img.onload = () => {{
        ctx.drawImage(img, 0, 0, canvas.width, canvas.height);
        canvas.toBlob(blob => {{
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url; a.download = '{self.title}.png'; a.click();
            URL.revokeObjectURL(url);
        }});
    }};
    img.src = 'data:image/svg+xml;base64,' + btoa(unescape(encodeURIComponent(svgData)));
}}

// 初始适应
setTimeout(fitAll, 100);
</script>
</body>
</html>'''
        return html

    def save_html(self, filepath, animated=True):
        """保存交互式 HTML 文件"""
        html = self.generate_html(animated=animated)
        os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html)
        print(f"交互式 HTML 已保存: {filepath}")
        return filepath

    # ---------- PNG 图标绘制（匹配 SVG 图标风格） ----------

    @staticmethod
    def _draw_png_icon_hexagon(draw, cx, cy, r, color='white'):
        """六边形图标（物质利益）"""
        pts = []
        for i in range(6):
            angle = math.pi / 6 + i * math.pi / 3
            pts.append((cx + r * math.cos(angle), cy + r * math.sin(angle)))
        draw.polygon(pts, fill=None, outline=color, width=2)

    @staticmethod
    def _draw_png_icon_shield(draw, cx, cy, r, color='white'):
        """盾牌图标（安全利益）"""
        w, h = r * 0.9, r * 1.1
        pts = [
            (cx, cy - h),                          # 顶点
            (cx + w, cy - h * 0.4),                 # 右上
            (cx + w * 0.7, cy + h * 0.5),           # 右下
            (cx, cy + h * 1.1),                     # 底部尖端
            (cx - w * 0.7, cy + h * 0.5),           # 左下
            (cx - w, cy - h * 0.4),                 # 左上
        ]
        draw.polygon(pts, fill=None, outline=color, width=2)
        # 盾牌中线
        draw.line([(cx, cy - h * 0.7), (cx, cy + h * 0.7)], fill=color, width=2)

    @staticmethod
    def _draw_png_icon_balance(draw, cx, cy, r, color='white'):
        """天平图标（政治利益）"""
        w = r * 0.9
        # 中柱
        draw.line([(cx, cy - r), (cx, cy + r * 0.6)], fill=color, width=3)
        # 底座
        draw.line([(cx - w * 0.6, cy + r * 0.6), (cx + w * 0.6, cy + r * 0.6)], fill=color, width=3)
        # 横梁
        draw.line([(cx - w, cy - r * 0.4), (cx + w, cy - r * 0.6)], fill=color, width=2)
        # 左盘（弧线近似）
        pts_left = []
        for i in range(9):
            t = i / 8
            x = cx - w * 1.1 + t * w
            y_off = (cy - r * 0.2) + (4 * t * (1 - t)) * r * 0.3
            pts_left.append((x, y_off))
        if len(pts_left) >= 2:
            draw.line(pts_left, fill=color, width=2)
        # 右盘
        pts_right = []
        for i in range(9):
            t = i / 8
            x = cx + w * 0.1 + t * w
            y_off = (cy - r * 0.4) + (4 * t * (1 - t)) * r * 0.3
            pts_right.append((x, y_off))
        if len(pts_right) >= 2:
            draw.line(pts_right, fill=color, width=2)

    @staticmethod
    def _draw_png_icon_masks(draw, cx, cy, r, color='white'):
        """双面具图标（身份与文化利益）"""
        r1 = r * 0.5
        offset = r * 0.4
        # 左面具（笑脸）
        draw.ellipse([cx - offset - r1, cy - r1, cx - offset + r1, cy + r1],
                     fill=None, outline=color, width=2)
        # 左面具微笑弧线
        smile_pts = []
        for i in range(7):
            t = i / 6
            x = cx - offset - r1 * 0.5 + t * r1
            y = cy + r1 * 0.1 + (4 * t * (1 - t)) * r1 * 0.3
            smile_pts.append((x, y))
        if len(smile_pts) >= 2:
            draw.line(smile_pts, fill=color, width=1)
        # 左面具眼睛
        draw.line([(cx - offset - r1 * 0.35, cy - r1 * 0.25),
                   (cx - offset - r1 * 0.05, cy - r1 * 0.25)], fill=color, width=2)
        draw.line([(cx - offset + r1 * 0.05, cy - r1 * 0.25),
                   (cx - offset + r1 * 0.35, cy - r1 * 0.25)], fill=color, width=2)
        # 右面具（哭脸）
        draw.ellipse([cx + offset - r1, cy - r1, cx + offset + r1, cy + r1],
                     fill=None, outline=color, width=2)
        # 右面具哭脸弧线
        cry_pts = []
        for i in range(7):
            t = i / 6
            x = cx + offset - r1 * 0.5 + t * r1
            y = cy + r1 * 0.4 - (4 * t * (1 - t)) * r1 * 0.25
            cry_pts.append((x, y))
        if len(cry_pts) >= 2:
            draw.line(cry_pts, fill=color, width=1)
        # 右面具眼睛
        draw.line([(cx + offset - r1 * 0.35, cy - r1 * 0.25),
                   (cx + offset - r1 * 0.05, cy - r1 * 0.25)], fill=color, width=2)
        draw.line([(cx + offset + r1 * 0.05, cy - r1 * 0.25),
                   (cx + offset + r1 * 0.35, cy - r1 * 0.25)], fill=color, width=2)

    @staticmethod
    def _draw_png_icon_pillar(draw, cx, cy, r, color='white'):
        """柱子图标（制度性与未来利益）"""
        pw = r * 0.5
        # 柱身
        draw.rectangle([cx - pw, cy - r * 0.6, cx + pw, cy + r * 0.6],
                       fill=None, outline=color, width=2)
        # 顶梁
        draw.line([(cx - pw * 1.4, cy - r * 0.6), (cx + pw * 1.4, cy - r * 0.6)],
                  fill=color, width=3)
        # 底座
        draw.line([(cx - pw * 1.4, cy + r * 0.6), (cx + pw * 1.4, cy + r * 0.6)],
                  fill=color, width=3)
        # 柱身竖线装饰
        draw.line([(cx, cy - r * 0.5), (cx, cy + r * 0.5)], fill=color, width=1)

    @staticmethod
    def _draw_png_icon_globe(draw, cx, cy, r, color='white'):
        """地球图标（公共利益）"""
        gr = r * 0.85
        # 外圆
        draw.ellipse([cx - gr, cy - gr, cx + gr, cy + gr],
                     fill=None, outline=color, width=2)
        # 中线（赤道）
        draw.line([(cx - gr, cy), (cx + gr, cy)], fill=color, width=1)
        # 经线（竖椭圆）
        erx = gr * 0.4
        draw.ellipse([cx - erx, cy - gr, cx + erx, cy + gr],
                     fill=None, outline=color, width=1)
        # 北半球弧线
        pts_n = []
        for i in range(9):
            t = i / 8
            x = cx - gr * 0.75 + t * gr * 1.5
            y = cy - gr * 0.45 - (4 * t * (1 - t)) * gr * 0.15
            pts_n.append((x, y))
        if len(pts_n) >= 2:
            draw.line(pts_n, fill=color, width=1)
        # 南半球弧线
        pts_s = []
        for i in range(9):
            t = i / 8
            x = cx - gr * 0.75 + t * gr * 1.5
            y = cy + gr * 0.45 + (4 * t * (1 - t)) * gr * 0.15
            pts_s.append((x, y))
        if len(pts_s) >= 2:
            draw.line(pts_s, fill=color, width=1)

    _PNG_ICON_DRAWERS = {
        "hexagon": _draw_png_icon_hexagon,
        "shield": _draw_png_icon_shield,
        "balance": _draw_png_icon_balance,
        "masks": _draw_png_icon_masks,
        "pillar": _draw_png_icon_pillar,
        "globe": _draw_png_icon_globe,
    }

    # ---------- PNG 静态图生成（使用 Pillow） ----------

    def generate_png(self, width=None, height=None, dpi=200):
        """
        使用 Pillow 生成 PNG 静态图片。
        用于 Word/PDF 报告嵌入。
        """
        if not HAS_PILLOW:
            raise ImportError("Pillow 未安装，无法生成 PNG")

        w = width or self.width
        h = height or self.height
        scale = w / self.width  # 缩放系数
        legend_h = int(self.LEGEND_HEIGHT * scale)
        total_h = int(h + legend_h)
        img = Image.new('RGBA', (w, total_h), '#FAFBFC')
        draw = ImageDraw.Draw(img)
        positions = self._compute_positions()
        type_config = INTEREST_TYPES
        flow_config = FLOW_TYPES
        R = self.NODE_RADIUS * scale  # 缩放

        # 尝试加载中文字体（增强字体搜索）
        font_paths = [
            "C:\\Windows\\Fonts\\msyh.ttc",
            "C:\\Windows\\Fonts\\msyhbd.ttc",
            "C:\\Windows\\Fonts\\simhei.ttf",
            "C:\\Windows\\Fonts\\simsun.ttc",
            "C:\\Windows\\Fonts\\simkai.ttf",
            "C:\\Windows\\Fonts\\simfang.ttf",
            "C:\\Windows\\Fonts\\msjh.ttc",
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/System/Library/Fonts/PingFang.ttc",
        ]
        font_large = font_small = font_tiny = None
        for fp in font_paths:
            if os.path.exists(fp):
                try:
                    font_large = ImageFont.truetype(fp, int(22 * scale))  # v6.1.1: 16→22 提升标题清晰度
                    font_small = ImageFont.truetype(fp, int(17 * scale))  # v6.1.1: 12→17 提升节点标签清晰度
                    font_tiny = ImageFont.truetype(fp, int(14 * scale))   # v6.1.1: 10→14 提升边标签/图例清晰度
                except Exception:
                    pass
                break
        if not font_large:
            font_large = ImageFont.load_default()
            font_small = font_large
            font_tiny = font_large

        # 辅助函数：居中绘制文字
        def draw_centered_text(xy, text, font, fill, anchor="mm"):
            """使用 textbbox 精确居中绘制文字"""
            bbox = draw.textbbox((0, 0), text, font=font)
            tw = bbox[2] - bbox[0]
            th = bbox[3] - bbox[1]
            x, y = xy
            draw.text((x - tw / 2, y - th / 2), text, fill=fill, font=font)

        # 绘制边
        for edge in self.edges:
            src_pos = positions.get(edge['source'])
            tgt_pos = positions.get(edge['target'])
            if not src_pos or not tgt_pos:
                continue
            ft = flow_config[edge['type']]
            sx, sy = src_pos[0] * scale, src_pos[1] * scale
            tx, ty = tgt_pos[0] * scale, tgt_pos[1] * scale

            dx, dy = tx - sx, ty - sy
            dist = max(math.sqrt(dx * dx + dy * dy), 1)
            ux, uy = dx / dist, dy / dist

            if edge['source'] != edge['target']:
                start_x = sx + ux * R
                start_y = sy + uy * R
                end_x = tx - ux * R
                end_y = ty - uy * R
            else:
                continue  # 自环在 Pillow 中不好画

            # 线条颜色
            color = ft['color']
            line_width = max(1, int(ft['width'] * scale))

            # 虚线
            dash_pattern = ft['dash']
            if dash_pattern == "none":
                draw.line([(start_x, start_y), (end_x, end_y)],
                         fill=color, width=line_width)
            else:
                parts = [int(p) * scale for p in dash_pattern.split(',')]
                d = 0
                drawing = True
                while d < dist - R * 2:
                    seg = parts[0] if drawing else parts[1] if len(parts) > 1 else parts[0]
                    d1 = d
                    d2 = min(d + seg, dist - R * 2)
                    if drawing:
                        x1 = start_x + ux * d1
                        y1 = start_y + uy * d1
                        x2 = start_x + ux * d2
                        y2 = start_y + uy * d2
                        draw.line([(x1, y1), (x2, y2)], fill=color, width=line_width)
                    d += seg
                    drawing = not drawing

            # 箭头
            arrow_size = ft['arrow_size'] * scale
            ax, ay = end_x, end_y
            p1x = ax - ux * arrow_size + uy * arrow_size * 0.4
            p1y = ay - uy * arrow_size - ux * arrow_size * 0.4
            p2x = ax - ux * arrow_size - uy * arrow_size * 0.4
            p2y = ay - uy * arrow_size + ux * arrow_size * 0.4
            draw.polygon([(ax, ay), (p1x, p1y), (p2x, p2y)], fill=color)

            # 标签
            if edge['label']:
                mid_x = (start_x + end_x) / 2
                mid_y = (start_y + end_y) / 2
                # 标签背景
                bbox = draw.textbbox((0, 0), edge['label'], font=font_small)
                tw = bbox[2] - bbox[0]
                th = bbox[3] - bbox[1]
                pad = 2
                draw.rectangle(
                    [mid_x - tw/2 - pad, mid_y - th/2 - 10*scale - pad,
                     mid_x + tw/2 + pad, mid_y + th/2 - 10*scale + pad],
                    fill='#FAFBFC', outline=None
                )
                draw_centered_text((mid_x, mid_y - 10 * scale), edge['label'], font_small, color)

        # 绘制节点
        for node in self.nodes:
            pos = positions.get(node['id'])
            if not pos:
                continue
            nx, ny = pos[0] * scale, pos[1] * scale
            tc = type_config[node['type']]

            # 光晕
            for i in range(3):
                r = R + 3 - i
                draw.ellipse(
                    [nx - r, ny - r, nx + r, ny + r],
                    fill=None, outline=tc['color'], width=1
                )

            # 背景圆
            draw.ellipse(
                [nx - R, ny - R, nx + R, ny + R],
                fill=tc['bg_color'], outline=tc['border_color'], width=max(1, int(2.5 * scale))
            )

            # 几何图标（节点内居中，使用边框色保证在浅背景上清晰可见）
            icon_r = R * 0.55
            icon_drawer = self._PNG_ICON_DRAWERS.get(tc['icon'])
            if icon_drawer and icon_r > 4:
                icon_drawer(draw, nx, ny, icon_r, color=tc['border_color'])

            # 类型短标签（图标下方，节点内底部）
            draw_centered_text((nx, ny + R * 0.7), tc['short_label'], font_tiny, tc['border_color'])

            # 名称标签（节点外下方，精确居中）
            draw_centered_text((nx, ny + R + 14 * scale), node['label'], font_small, '#333')

        # 标题（精确居中）
        draw_centered_text((w / 2, 20 * scale), self.title, font_large, '#1B3A5C')

        # 图例
        self._draw_png_legend(draw, type_config, flow_config, w, h, scale, font_tiny)

        return img

    def _draw_png_legend(self, draw, type_config, flow_config, w, h, scale, font):
        """在 PNG 底部绘制横向图例条带（与节点区域物理隔离）"""
        legend_y = h + 10 * scale  # 图例区在节点区下方 10px
        lh = self.LEGEND_HEIGHT * scale

        # 分隔线（虚线）
        dash_len = 6 * scale
        gap_len = 3 * scale
        y_sep = h
        x = 15 * scale
        while x < w - 15 * scale:
            draw.line([(x, y_sep), (min(x + dash_len, w - 15 * scale), y_sep)],
                       fill='#D0D0D0', width=1)
            x += dash_len + gap_len

        # --- 第一行：利益类型（6个水平排列） ---
        draw.text((15 * scale, legend_y + 2 * scale), "利益类型", fill='#333', font=font)
        type_items = list(type_config.items())
        for i, (tkey, tconf) in enumerate(type_items):
            start_x = 100 * scale
            spacing = 155 * scale  # v6.1.1: 125→155 容纳完整中文名称
            cx = start_x + i * spacing
            cy = legend_y - 4 * scale
            r = 6 * scale
            draw.ellipse([cx - r, cy - r, cx + r, cy + r],
                          fill=tconf['color'], outline=tconf['border_color'], width=1)
            draw.text((cx + 10 * scale, cy - 6 * scale),
                       f"{tconf['short_label']} {tconf['name']}", fill='#555', font=font)

        # --- 第二行：利益动线（4个水平排列） ---
        ly2 = legend_y + 26 * scale
        draw.text((15 * scale, ly2), "利益动线", fill='#333', font=font)
        flow_items = list(flow_config.items())
        for i, (fkey, fconf) in enumerate(flow_items):
            start_x = 100 * scale
            spacing = 190 * scale
            lx = start_x + i * spacing
            ly = ly2 - 2 * scale
            draw.line([(lx, ly + 4 * scale), (lx + 30 * scale, ly + 4 * scale)],
                       fill=fconf['color'], width=2)
            draw.text((lx + 35 * scale, ly), fconf['name'], fill='#555', font=font)

    def save_png(self, filepath, width=None, height=None, dpi=200):
        """保存 PNG 文件"""
        img = self.generate_png(width, height, dpi)
        os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
        img.save(filepath, 'PNG', dpi=(dpi, dpi))
        print(f"PNG 已保存: {filepath}")
        return filepath

    # ---------- Base64 编码 ----------

    def generate_base64_png(self, width=None, height=None, dpi=200):
        """生成 Base64 编码的 PNG 图片"""
        img = self.generate_png(width, height, dpi)
        buffer = BytesIO()
        img.save(buffer, 'PNG')
        b64 = base64.b64encode(buffer.getvalue()).decode('ascii')
        return f"data:image/png;base64,{b64}"

    def generate_base64_svg(self, animated=False):
        """生成 Base64 编码的 SVG"""
        svg_str = self.generate_svg(animated=animated)
        b64 = base64.b64encode(svg_str.encode('utf-8')).decode('ascii')
        return f"data:image/svg+xml;base64,{b64}"

    # ---------- Word 文档嵌入 ----------

    def embed_to_word(self, doc, width_inches=5.5, figure_number=None):
        """
        将网络图作为 PNG 图片嵌入到 python-docx 的 Document 对象中。

        Args:
            doc: python-docx Document 对象
            width_inches: 图片显示宽度（英寸）
            figure_number: 图片序号（如 1, 2, 3），用于生成"图X：标题"格式的题注
        """
        if not HAS_DOCX:
            raise ImportError("python-docx 未安装")

        # 生成高分辨率 PNG
        img = self.generate_png(width=int(self.width * 2), height=int(self.height * 2))
        buffer = BytesIO()
        img.save(buffer, 'PNG', dpi=(200, 200))
        buffer.seek(0)

        from docx.shared import Inches
        doc.add_picture(buffer, width=Inches(width_inches))

        # 图片居中
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

        # 图片下方题注（含图号）
        from docx.shared import Pt, RGBColor
        cap = doc.add_paragraph()
        cap.alignment = 1
        cap_text = f"图{figure_number}：{self.title}" if figure_number else f"图：{self.title}"
        run = cap.add_run(cap_text)
        from docx.oxml.ns import qn
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

    # ---------- 便捷方法 ----------

    @classmethod
    def from_interest_analysis(cls, analysis_data, title="利益关系网络"):
        """
        从结构化利益分析数据快速构建可视化。

        Args:
            analysis_data: dict，格式如下：
                {
                    "nodes": [
                        {"id": "A", "label": "地方政府", "type": "political", "detail": "..."},
                        ...
                    ],
                    "edges": [
                        {"source": "A", "target": "B", "label": "审批权→利润", "type": "economic"},
                        ...
                    ]
                }
            title: 网络图标题
        """
        viz = cls(title=title)
        for node in analysis_data.get('nodes', []):
            viz.add_node(
                node['id'], node.get('label', node['id']),
                node.get('type', 'material'),
                detail=node.get('detail', ''),
                group=node.get('group', '')
            )
        for edge in analysis_data.get('edges', []):
            viz.add_edge(
                edge['source'], edge['target'],
                label=edge.get('label', ''),
                flow_type=edge.get('type', 'economic'),
                detail=edge.get('detail', ''),
                bidirectional=edge.get('bidirectional', False)
            )
        return viz

    def summary(self):
        """返回网络摘要信息"""
        type_counts = {}
        for n in self.nodes:
            t = n['type']
            type_counts[t] = type_counts.get(t, 0) + 1
        flow_counts = {}
        for e in self.edges:
            t = e['type']
            flow_counts[t] = flow_counts.get(t, 0) + 1
        return {
            "title": self.title,
            "nodes": len(self.nodes),
            "edges": len(self.edges),
            "node_types": type_counts,
            "flow_types": flow_counts,
            "layout": self.layout_algo,
        }


# ============================================================
# 独立运行入口
# ============================================================

if __name__ == "__main__":
    # 演示：大同订婚案利益关系网络
    viz = InterestNetworkViz(
        title="大同订婚强奸案 · 利益关系网络",
        width=960,
        height=720,
        layout="force"
    )

    # 利益主体节点
    viz.add_node("court", "阳高县法院", "political", detail="审判机关，掌握司法裁判权与量刑尺度，受上级法院和舆论双重压力")
    viz.add_node("prosecutor", "检察院", "political", detail="公诉方，代表国家追究刑事责任，受考核指标影响")
    viz.add_node("bride", "女方（席某）", "security", detail="强奸指控方，主张人身安全与性自主权受侵害，同时涉及彩礼利益")
    viz.add_node("groom", "男方（甄某）", "material", detail="被指控方，面临刑事处罚，主张婚约财产返还，安全利益受威胁")
    viz.add_node("bride_family", "女方家属", "identity_culture", detail="彩礼谈判主体，维护家庭面子与经济利益，推动刑事追诉")
    viz.add_node("groom_family", "男方家属", "identity_culture", detail="维权主体，主张婚约被骗，寻求司法公正与财产返还")
    viz.add_node("public", "网络舆论", "public", detail="放大器与审判场，性别对立叙事的竞技场，司法公信力的试金石")
    viz.add_node("media", "媒体", "identity_culture", detail="叙事生产者，选择性放大冲突点，流量驱动下的议程设置")
    viz.add_node("law_sys", "司法系统", "institutional_future", detail="制度框架，强奸罪认定的证据标准与婚约财产规则的冲突点")

    # 利益动线
    viz.add_edge("bride", "prosecutor", "性自主权侵害主张", "legal")
    viz.add_edge("prosecutor", "court", "公诉请求", "power")
    viz.add_edge("bride_family", "bride", "推动刑事追诉", "power")
    viz.add_edge("groom_family", "groom", "维权支持", "cultural")
    viz.add_edge("bride_family", "groom_family", "彩礼返还争议", "economic")
    viz.add_edge("court", "groom", "刑事判决", "legal")
    viz.add_edge("court", "bride", "性自主权保护", "legal")
    viz.add_edge("public", "court", "舆论压力", "cultural")
    viz.add_edge("media", "public", "叙事扩散", "cultural")
    viz.add_edge("law_sys", "court", "证据标准约束", "legal")
    viz.add_edge("groom", "bride_family", "彩礼返还诉求", "economic")
    viz.add_edge("public", "law_sys", "公信力质疑", "power")

    # 输出
    base_dir = os.path.dirname(os.path.abspath(__file__))
    output_dir = os.path.join(base_dir, "报告", "viz_demo")
    os.makedirs(output_dir, exist_ok=True)

    viz.save_html(os.path.join(output_dir, "利益关系网络.html"))
    viz.save_png(os.path.join(output_dir, "利益关系网络.png"), width=1440, height=1080)

    print(f"\n网络摘要: {json.dumps(viz.summary(), ensure_ascii=False, indent=2)}")
